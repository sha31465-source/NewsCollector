"""
资讯采集系统 - 全功能独立版本
====================================
版本: V3.4 Standalone
说明: 这是一个单文件独立版本，所有代码集中在一个文件中
只需 classification_rules.yaml 配置文件即可运行

作者: Claude Code (Sonnet 4.5)
日期: 2026-04-04
"""

import customtkinter as ctk
import threading
import sys
import os
import traceback
import sqlite3
import hashlib
import requests
import feedparser
import yaml
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import List, Dict, Optional
import re
import time
import ssl
import urllib3
from docx import Document
from tkinter import messagebox
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================
# CSV导出模块
# ============================================
try:
    import csv_exporter
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False
    print("[警告] CSV导出模块未找到，将仅生成TXT报告")

# 禁用SSL警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ============================================
# 源代码提取功能
# ============================================

def extract_source_if_needed():
    """
    如果源文件不存在，从PyInstaller临时目录中提取源代码
    这样即使源文件被删除，也可以从exe中恢复
    """
    # 如果是PyInstaller打包的exe
    if getattr(sys, 'frozen', False):
        # 当前脚本目录
        script_dir = os.path.dirname(os.path.abspath(sys.executable))

        # 源文件路径
        source_py = os.path.join(script_dir, 'main_gui_standalone.py')
        source_spec = os.path.join(script_dir, 'NewsCollector_V34.spec')

        # 如果源文件不存在，尝试提取
        if not os.path.exists(source_py):
            try:
                import shutil

                # PyInstaller临时目录
                if '_MEIPASS' in os.environ:
                    temp_dir = os.environ['_MEIPASS']

                    # 提取源文件
                    temp_py = os.path.join(temp_dir, 'source', 'main_gui_standalone.py')
                    temp_spec = os.path.join(temp_dir, 'source', 'NewsCollector_V34.spec')

                    if os.path.exists(temp_py):
                        shutil.copy2(temp_py, source_py)
                        print(f"[源代码] 已提取源文件到: {source_py}")

                    if os.path.exists(temp_spec):
                        shutil.copy2(temp_spec, source_spec)
                        print(f"[源代码] 已提取配置文件到: {source_spec}")

            except Exception as e:
                print(f"[源代码] 提取失败: {e}")

# ============================================
# 文件路径配置
# ============================================

# 基础输出路径
BASE_OUTPUT_PATH = r"M:\claude\生成文件"

# 获取当前日期的输出目录
def get_output_dirs():
    """获取各类型文件的输出目录"""
    date_str = datetime.now().strftime('%Y-%m-%d')
    base_dir = os.path.join(BASE_OUTPUT_PATH, date_str)

    dirs = {
        'base': base_dir,
        'scripts': os.path.join(base_dir, '脚本'),
        'logs': os.path.join(base_dir, '日志'),
        'docs': os.path.join(base_dir, '文档'),
        'images': os.path.join(base_dir, '图片'),
        'other': os.path.join(base_dir, '其他')
    }

    # 创建所有目录
    for dir_path in dirs.values():
        os.makedirs(dir_path, exist_ok=True)

    return dirs

# ============================================
# RSS源配置
# ============================================

MEDICAL_RSS_SOURCES = [
    'https://www.sciencedaily.com/rss/health_medicine/lung_cancer.xml',
    'https://www.sciencedaily.com/rss/health_medicine/cancer.xml',
    'https://www.sciencedaily.com/rss/health_medicine.xml',
    'https://www.sciencedaily.com/rss/top/health.xml',
    'https://www.frontiersin.org/journals/oncology/rss',
    'https://www.nature.com/nrclinonc/rss/current',
    'https://www.nature.com/natcancer/rss/current',
    'https://www.medscape.com/viewarticle/rss',
    'https://www.nature.com/nutd/rss/current',
    'https://www.sciencedaily.com/rss/health_medicine/nutrition.xml',
    'https://www.frontiersin.org/journals/nutrition/rss',
    # 移除失效的RSS源
    # 'https://www.sciencedaily.com/rss/health_medicine/stem_cells.xml',  # SSL错误
    'https://www.sciencedaily.com/rss/health_medicine/immune_system.xml',
    # 'https://www.sciencedaily.com/rss/health_medicine/genetics.xml',  # 404
    'https://www.frontiersin.org/journals/public-health/rss',
    # 'https://www.nature.com/articles?type=research,news&subject=medical-research',  # 404
    # 'https://www.sciencedaily.com/rss/health_medicine/diagnostics.xml',  # 404
    # 'https://www.sciencedaily.com/rss/health_medicine/pharmaceuticals.xml',  # SSL错误
    'https://www.frontiersin.org/journals/surgery/rss',
    # 新增替代源
    'https://www.nature.com/nm/rss',  # Nature Medicine
    'https://www.nature.com/ng/rss',  # Nature Genetics
]

AI_TECH_RSS_SOURCES = [
    'https://www.technologyreview.com/feed/',
    'https://techcrunch.com/feed/',
    'https://feeds.arstechnica.com/arstechnica/index',
    'https://www.wired.com/feed/rss',
    'https://openai.com/news/rss.xml',
    'https://blog.google/technology/ai/rss/',
    'https://www.sciencedaily.com/rss/computers_math.xml',
    'https://anthropic.com/news/rss',
    # 移除失效的RSS源
    # 'https://www.anthropic.com/index/rss',  # 308重定向
    # 更新为新的The Verge AI RSS地址
    'https://www.theverge.com/rss/ai-artificial-intelligence/index.xml',
    'https://www.technologyreview.com/topic/artificial-intelligence/feed/',
    'https://arstechnica.com/tag/ai/feed/',
    # 'https://www.wired.com/feed/category/ai/latest/rss',  # 404
]

BREAKING_NEWS_SOURCES = [
    'https://www.businessinsider.com/rss',
    'https://jalopnik.com/rss',
    'https://www.torquenews.com/rss.xml',
    'https://electrek.co/feed/',
    'https://www.ithome.com/rss/',
    'https://feeds.bbci.co.uk/news/science_and_environment/rss.xml',
    'https://feeds.npr.org/1001/rss.xml',
    'https://feeds.bbci.co.uk/news/world/rss.xml',
    # 移除失效的RSS源
    # 'https://www.reuters.com/news/technology',  # 原地址失效
    # 'http://feeds.reuters.com/reuters/technologyNews',  # 400错误 - Reuters RSS已不可用
    'https://www.engadget.com/rss.xml',
    'https://www.cnet.com/rss/news/',
    'https://feeds.macrumors.com/MacRumors-All',
    'https://9to5mac.com/feed/',
    # 'https://www.androidpolice.com/feed/',  # 连接重置
    'https://www.theverge.com/rss/index.xml',
    'https://techcrunch.com/category/apps/feed/',
    'https://www.zdnet.com/news/rss.xml',
    # 新增替代源
    'https://www.cnbc.com/id/100003114/device/rss/rss.html',  # CNBC Technology
]

# 影视娱乐RSS源（新增）
MOVIE_ENTERTAINMENT_RSS_SOURCES = [
    # 移除失效的RSS源
    # 'https://www.rottentomatoes.com/trending/rss/',  # 404 - 官方已不再提供RSS
    'https://www.hollywoodreporter.com/feed/',  # The Hollywood Reporter - 侧重于行业新闻、制作动态、奖项预测
    # 'https://www.cinemablend.com/rss/news/movie',  # 404 - 官方已不再提供RSS
    # 新增替代源
    'https://www.variety.com/feed/',  # Variety - 权威的娱乐产业新闻
]

# 游戏领域RSS源（新增）
GAMING_RSS_SOURCES = [
    # 境内中文源
    'https://www.gcores.com/rss',  # 机核网 - 深度内容居多，适合提取长文观点
    # 【移除】其乐RSS源description字段为空，导致无法采集
    # 'https://keylol.com/forum.php?mod=rss',

    # 国外英文源
    'https://www.eurogamer.net/feed/news',  # Eurogamer - 权威的欧洲游戏媒体，技术分析和行业动态
    'https://www.pcgamer.com/rss/',  # PC Gamer - 专注PC平台，硬件与软件动态平衡
]

# 合并所有RSS源
ALL_RSS_SOURCES = (
    MEDICAL_RSS_SOURCES +
    AI_TECH_RSS_SOURCES +
    BREAKING_NEWS_SOURCES +
    MOVIE_ENTERTAINMENT_RSS_SOURCES +
    GAMING_RSS_SOURCES
)

# ============================================
# 数据库操作类
# ============================================

class SimpleDatabase:
    """简化的数据库操作类"""

    def __init__(self, db_path: str):
        self.db_path = db_path
        self.conn = None
        self._init_db()

    def _init_db(self):
        """初始化数据库表"""
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        cursor = self.conn.cursor()

        # 检查表是否存在
        cursor.execute('SELECT name FROM sqlite_master WHERE type="table"')
        tables = [row[0] for row in cursor.fetchall()]

        # 如果表不存在，创建新表结构
        if 'news_items' not in tables:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS news_items (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT,
                    content TEXT,
                    title_original TEXT,
                    content_original TEXT,
                    link TEXT UNIQUE,
                    pub_date TEXT,
                    source TEXT,
                    category_key TEXT,
                    category_tag TEXT,
                    status TEXT DEFAULT 'pending',
                    added_at TEXT,
                    translated_at TEXT,
                    processed_at TEXT,
                    analysis TEXT,
                    is_processed INTEGER DEFAULT 0
                )
            ''')
            self.conn.commit()
        else:
            # 表已存在，检查并添加缺失的字段
            cursor.execute('PRAGMA table_info(news_items)')
            columns = [row[1] for row in cursor.fetchall()]

            # 添加 analysis 字段（如果不存在）
            if 'analysis' not in columns:
                cursor.execute('ALTER TABLE news_items ADD COLUMN analysis TEXT')
                self.conn.commit()

            # 添加 is_processed 字段（如果不存在）
            cursor.execute('PRAGMA table_info(news_items)')
            columns = [row[1] for row in cursor.fetchall()]
            if 'is_processed' not in columns:
                cursor.execute('ALTER TABLE news_items ADD COLUMN is_processed INTEGER DEFAULT 0')
                self.conn.commit()
                print("[数据库] 已添加 is_processed 字段")

            # 添加 summary 字段（如果不存在）
            cursor.execute('PRAGMA table_info(news_items)')
            columns = [row[1] for row in cursor.fetchall()]
            if 'summary' not in columns:
                cursor.execute('ALTER TABLE news_items ADD COLUMN summary TEXT')
                self.conn.commit()
                print("[数据库] 已添加 summary 字段")

        # 创建 rss_sources 表（如果不存在）
        if 'rss_sources' not in tables:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS rss_sources (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    url TEXT NOT NULL UNIQUE,
                    name TEXT,
                    category TEXT,
                    source_type TEXT,
                    is_enabled INTEGER DEFAULT 1,
                    created_at TEXT,
                    last_checked_at TEXT
                )
            ''')
            self.conn.commit()
            print("[数据库] 已创建 rss_sources 表")

            # 插入默认RSS源
            self._insert_default_rss_sources(cursor)
        else:
            # 表已存在，检查是否有数据
            cursor.execute('SELECT COUNT(*) FROM rss_sources')
            count = cursor.fetchone()[0]
            if count == 0:
                print("[数据库] rss_sources表为空，插入默认数据")
                self._insert_default_rss_sources(cursor)

        # 创建 custom_keywords 表（如果不存在）
        if 'custom_keywords' not in tables:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS custom_keywords (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    keyword TEXT NOT NULL UNIQUE,
                    description TEXT,
                    is_enabled INTEGER DEFAULT 1,
                    created_at TEXT,
                    category TEXT DEFAULT 'default'
                )
            ''')
            self.conn.commit()
            print("[数据库] 已创建 custom_keywords 表")

    def _insert_default_rss_sources(self, cursor):
        """插入默认RSS源到数据库（仅国内可访问，无需VPN）"""
        from datetime import datetime
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # 默认RSS源列表（精选国内可访问源，无需VPN）
        default_sources = [
            # ========== 科技类（国内可访问） ==========
            ('https://www.ithome.com/rss/', 'IT之家', 'tech', 'tech'),
            ('https://www.sciencedaily.com/rss/computers_math.xml', 'Science Daily - 计算机', 'tech', 'tech'),
            ('https://www.sciencedaily.com/rss/top/tech.xml', 'Science Daily - 科技头条', 'tech', 'tech'),
            ('https://feeds.bbci.co.uk/news/technology/rss.xml', 'BBC - 科技', 'tech', 'tech'),

            # ========== 医学类（国内可访问） ==========
            ('https://www.sciencedaily.com/rss/health_medicine.xml', 'Science Daily - 医学', 'medical', 'medical'),
            ('https://www.sciencedaily.com/rss/top/health.xml', 'Science Daily - 健康头条', 'medical', 'medical'),
            ('https://www.sciencedaily.com/rss/health_medicine/nutrition.xml', 'Science Daily - 营养', 'medical', 'medical'),
            ('https://www.sciencedaily.com/rss/health_medicine/immune_system.xml', 'Science Daily - 免疫', 'medical', 'medical'),

            # ========== 综合新闻类（国内可访问） ==========
            ('https://www.ithome.com/rss/', 'IT之家 - 综合', 'news', 'news'),
            ('https://feeds.bbci.co.uk/news/world/rss.xml', 'BBC - 世界新闻', 'news', 'news'),
            ('https://feeds.bbci.co.uk/news/science_and_environment/rss.xml', 'BBC - 科环', 'news', 'news'),
            ('https://www.engadget.com/rss.xml', 'Engadget', 'news', 'news'),
            ('https://www.cnet.com/rss/news/', 'CNET - 新闻', 'news', 'news'),
        ]

        # 插入默认源
        for url, name, category, source_type in default_sources:
            try:
                cursor.execute('''
                    INSERT OR IGNORE INTO rss_sources (url, name, category, source_type, is_enabled, created_at)
                    VALUES (?, ?, ?, ?, 1, ?)
                ''', (url, name, category, source_type, now))
            except Exception as e:
                print(f"[数据库] 插入RSS源失败 {url}: {e}")

        self.conn.commit()
        print(f"[数据库] 已插入 {len(default_sources)} 个默认RSS源")

    def is_duplicate(self, url: str) -> bool:
        """检查是否重复（基于URL）"""
        cursor = self.conn.cursor()
        cursor.execute('SELECT id FROM news_items WHERE link = ?', (url,))
        return cursor.fetchone() is not None

    def is_duplicate_by_title(self, title: str, threshold: float = 0.85) -> bool:
        """检查标题是否与已有资讯相似（使用混合相似度算法：编辑距离 + Jaccard）

        Args:
            title: 待检查的标题
            threshold: 相似度阈值（0-1），默认0.85表示85%相似

        增强：
        - 查重窗口从14天扩展到90天
        - 使用编辑距离算法（考虑字序和修改代价）
        - 极致清洗：去除括号内容、标点、空格
        - 包含关系检查：处理长短标题重复
        - 详细的日志记录
        """
        cursor = self.conn.cursor()

        # 【修改】从14天扩展到90天，拦截周期性新闻
        cursor.execute('''
            SELECT title, added_at FROM news_items
            WHERE DATE(added_at) >= DATE('now', '-90 days')
        ''')
        rows = cursor.fetchall()

        if not rows:
            return False

        # 检查是否有相似标题
        for row in rows:
            existing_title = row[0]
            added_at = row[1]

            # 【增强】使用混合相似度算法（编辑距离 + Jaccard）
            similarity = self._calculate_similarity_hybrid(title, existing_title)

            # 【新增】详细日志：显示使用的算法类型
            text_length = max(len(self._clean_title_extreme(title)),
                            len(self._clean_title_extreme(existing_title)))
            algorithm = "编辑距离" if text_length < 20 else "混合算法(编辑距离60%+Jaccard40%)"

            # 相似度接近阈值时记录日志（便于调试）
            if similarity >= threshold * 0.9:
                print(f"[查重预警] 算法: {algorithm} | 相似度: {similarity:.3f} | 阈值: {threshold:.3f}")
                print(f"  新标题: {title[:60]}...")
                print(f"  旧标题: {existing_title[:60]}...")
                print(f"  采集时间: {added_at}")

            if similarity >= threshold:
                print(f"[标题去重] 检测到重复！算法: {algorithm} | 相似度: {similarity:.3f}")
                print(f"  新标题: {title[:80]}...")
                print(f"  旧标题: {existing_title[:80]}...")
                return True

        return False

    def is_duplicate_by_content(self, content: str, threshold: float = 0.85) -> bool:
        """检查内容是否与已有资讯相似（使用优化的Jaccard相似度）

        Args:
            content: 待检查的内容
            threshold: 相似度阈值（0-1），默认0.85表示85%相似
        """
        cursor = self.conn.cursor()

        # 使用增强版清洗方法预处理内容（移除HTML实体、脚本、样式）
        cleaned_content = self._clean_html_pro(content)

        # 获取最近14天内的所有内容
        cursor.execute('''
            SELECT content FROM news_items
            WHERE DATE(added_at) >= DATE('now', '-14 days')
        ''')
        rows = cursor.fetchall()

        if not rows:
            return False

        # 检查是否有相似内容
        for row in rows:
            existing_content = row[0]
            if not existing_content:
                continue

            # 同样清洗已存在的内容
            cleaned_existing = self._clean_html_pro(existing_content)

            similarity = self._calculate_similarity(cleaned_content, cleaned_existing, is_content=True)
            if similarity >= threshold:
                return True

        return False

    def _clean_title(self, title: str) -> str:
        """清洗标题，移除repost、via等后缀词（增强版：支持前缀和更多模式）"""
        import re
        if not title:
            return title

        # 转小写
        title = title.lower()

        # 【增强】移除常见前缀词（自媒体常用的"标题党"前缀）
        prefixes = [
            r'^【[^】]*】\s*',           # 【xxx】标题
            r'^\[[^\]]*\]\s*',          # [xxx]标题
            r'^〔[^〕]*〕\s*',           # 〔xxx〕标题
            r'^\*\*[^*]*\*\*\s*',       # **xxx**标题
            r'^#+\s*',                  # Markdown 标题
            r'^\d+[\.\、]\s*',          # 1. 标题
            r'^紧急[:：]\s*',           # 紧急：标题
            r'^重大[:：]\s*',           # 重大：标题
            r'^突发[:：]\s*',           # 突发：标题
            r'^最新[:：]\s*',           # 最新：标题
            r'^重磅[:：]\s*',           # 重磅：标题
            r'^独家[:：]\s*',           # 独家：标题
            r'^首发[:：]\s*',           # 首发：标题
            r'^深度[:：]\s*',           # 深度：标题
            r'^详细[:：]\s*',           # 详细：标题
            r'^完整[:：]\s*',           # 完整：标题
            r'^原创[:：]\s*',           # 原创：标题
        ]

        for prefix in prefixes:
            title = re.sub(prefix, '', title, flags=re.IGNORECASE)

        # 移除常见的后缀词（保留原有逻辑并扩展）
        suffixes = [
            r'\s*[-–—|｜]\s*repost.*$',     # - repost xxx
            r'\s*[-–—|｜]\s*via.*$',        # - via xxx
            r'\s*[-–—|｜]\s*source.*$',     # - source xxx
            r'\s*[-–—|｜]\s*translated.*$', # - translated by xxx
            r'\s*[-–—|｜]\s*[（\(].*?[）\)]$',  # - (xxx)
            r'\s+repost.*$',              # 空格 + repost xxx
            r'\s+via.*$',                 # 空格 + via xxx
            r'\s+来源.*$',                # 空格 + 来源 xxx
            r'\s+来源[:：].*$',           # 空格 + 来源：xxx
            r'\s+转载.*$',                # 空格 + 转载xxx
            r'\s+转载自.*$',              # 空格 + 转载自xxx
            r'\s+更多.*$',                # 空格 + 更多xxx
        ]

        for suffix in suffixes:
            title = re.sub(suffix, '', title, flags=re.IGNORECASE)

        # 【增强】移除所有全角和半角分隔符（替换为空格）
        title = re.sub(r'[-–—|｜／／\{\}\[\]〈〉《》【】「」『』（）()\s]+', ' ', title)

        # 移除多余空格
        title = re.sub(r'\s+', ' ', title)

        return title.strip()

    def _clean_title_extreme(self, title: str) -> str:
        """极致清洗标题：去除括号内容、标点、空格，只保留核心文字

        用于编辑距离算法的预处理，比 _clean_title 更激进

        Examples:
            "【突发】重大AI突破！（视频）" → "重大ai突破"
            "[转载]OpenAI发布新模型 via Twitter" → "openai发布新模型"
        """
        import re
        if not title:
            return ""

        # 去除括号及其内容（包括中英文括号）
        # 例如: [转载], (图), 【视频】, （图片）
        text = re.sub(r'[\[\(【].*?[\]\)】]', '', title)

        # 仅保留中文字符、字母和数字
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)

        return text.lower()

    def _calculate_similarity_edit_distance(self, str1: str, str2: str) -> float:
        """基于归一化编辑距离（Levenshtein Distance）计算相似度

        使用 Python 内置的 difflib.SequenceMatcher，基于 Ratcliff/Obershelp 算法
        该算法考虑字序和修改代价，比 Jaccard 更适合检测"标题党"

        Args:
            str1: 字符串1
            str2: 字符串2

        Returns:
            相似度（0-1之间的浮点数）

        Examples:
            "AI医疗新突破" vs "重大突破：AI医疗新算法"
            Jaccard: 0.67 (只看字出现)
            Edit Distance: 0.78 (考虑字序)
        """
        from difflib import SequenceMatcher

        # 极致清洗
        s1 = self._clean_title_extreme(str1)
        s2 = self._clean_title_extreme(str2)

        if not s1 or not s2:
            return 0.0

        # 1. 检查包含关系（处理长短标题重复）
        if s1 in s2 or s2 in s1:
            return 1.0

        # 2. 计算序列相似度（基于动态规划的 Ratcliff/Obershelp 算法）
        # 这个算法类似于编辑距离，但 Python 实现优化过，性能更好
        matcher = SequenceMatcher(None, s1, s2)
        return matcher.ratio()

    def _calculate_similarity_hybrid(self, str1: str, str2: str) -> float:
        """混合相似度计算：结合 Jaccard 和编辑距离

        策略：
        - 短标题（<20字）：优先使用编辑距离（更精确）
        - 长标题（>=20字）：使用加权平均（综合两种算法）

        Args:
            str1: 字符串1
            str2: 字符串2

        Returns:
            综合相似度（0-1之间的浮点数）
        """
        import re

        # 预处理：转小写
        str1_lower = str1.lower()
        str2_lower = str2.lower()

        # 极致清洗后的版本（用于编辑距离）
        s1_extreme = self._clean_title_extreme(str1)
        s2_extreme = self._clean_title_extreme(str2)

        # 如果任一字符串为空，返回0
        if not s1_extreme or not s2_extreme:
            return 0.0

        # 计算编辑距离相似度
        edit_sim = self._calculate_similarity_edit_distance(str1, str2)

        # 计算Jaccard相似度（保留原有逻辑作为辅助）
        def tokenize(text):
            tokens = []
            chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
            tokens.extend(chinese_chars)
            english_words = re.findall(r'[a-zA-Z]+', text)
            tokens.extend([w.lower() for w in english_words])
            numbers = re.findall(r'\d+', text)
            tokens.extend(numbers)
            return tokens

        set1 = set(tokenize(s1_extreme))
        set2 = set(tokenize(s2_extreme))

        if not set1 or not set2:
            jaccard_sim = 0.0
        else:
            intersection = set1 & set2
            union = set1 | set2
            jaccard_sim = len(intersection) / len(union) if union else 0.0

        # 根据标题长度选择策略
        text_length = max(len(s1_extreme), len(s2_extreme))

        if text_length < 20:
            # 短标题：编辑距离更准确
            final_sim = edit_sim
        else:
            # 长标题：加权平均（编辑距离权重60%，Jaccard权重40%）
            final_sim = edit_sim * 0.6 + jaccard_sim * 0.4

        return final_sim

    def _preprocess_content(self, text: str) -> str:
        """预处理内容：移除HTML标签和特殊字符"""
        import re
        # 转小写
        text = text.lower()
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', ' ', text)
        # 移除特殊字符，保留中英文、数字和空格
        text = re.sub(r'[^\w\s\u4e00-\u9fff]', ' ', text)
        # 移除多余空格
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def save_item(self, item: dict) -> int:
        """保存资讯项"""
        cursor = self.conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO news_items (title, content, link, source, pub_date, category_key, added_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                item.get('title'),
                item.get('content'),
                item.get('url'),
                item.get('source'),
                item.get('pub_date') or item.get('published_date'),
                item.get('category', '未分类'),
                datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            ))
            self.conn.commit()
            return cursor.lastrowid
        except sqlite3.IntegrityError:
            return -1

    def get_raw_items(self, status: str = 'pending', limit: int = None) -> List[dict]:
        """获取未处理的资讯"""
        cursor = self.conn.cursor()
        query = 'SELECT * FROM news_items WHERE status = ?'
        params = [status]

        if limit:
            query += ' LIMIT ?'
            params.append(limit)

        cursor.execute(query, params)
        rows = cursor.fetchall()

        columns = [desc[0] for desc in cursor.description]
        return [dict(zip(columns, row)) for row in rows]

    def get_processed_items(self, limit: int = None, date: str = None, use_processed_time: bool = False) -> List[dict]:
        """获取已处理的资讯

        Args:
            limit: 限制数量
            date: 指定日期 (格式: 'YYYY-MM-DD')
            use_processed_time: 如果为True，按处理时间（processed_at）过滤；如果为False，按采集时间（added_at）过滤

        说明：
            使用 is_processed 字段判断是否已处理（优先）
            只返回 status = 'processed' 的内容（排除已生成报告的 'reported' 状态）
            兼容旧数据：如果 is_processed 为 NULL，则检查 analysis 字段
        """
        cursor = self.conn.cursor()
        # 使用 is_processed 字段判断（兼容旧数据）
        # 增加 status = 'processed' 的过滤条件，排除已生成报告的 'reported' 状态
        query = '''
            SELECT * FROM news_items
            WHERE (is_processed = 1
               OR (is_processed IS NULL AND analysis IS NOT NULL))
              AND status = 'processed'
        '''
        params = []

        # 如果指定了日期，根据 use_processed_time 参数决定按哪个时间字段过滤
        if date:
            if use_processed_time:
                # 按处理时间过滤（用于生成本次运行的报告）
                # 使用 "<=" 可以包含该日期及之前所有未报告的资讯，避免遗漏
                query += ' AND DATE(processed_at) <= ?'
            else:
                # 按采集时间过滤（原逻辑）
                query += ' AND DATE(added_at) = ?'
            params.append(date)

        if limit:
            query += ' LIMIT ?'
            params.append(limit)

        cursor.execute(query, params)
        rows = cursor.fetchall()

        columns = [desc[0] for desc in cursor.description]
        return [dict(zip(columns, row)) for row in rows]

    def update_item(self, item_id: int, **kwargs):
        """更新资讯项"""
        cursor = self.conn.cursor()
        set_clause = ', '.join([f'{k} = ?' for k in kwargs.keys()])
        query = f'UPDATE news_items SET {set_clause} WHERE id = ?'
        values = list(kwargs.values()) + [item_id]
        cursor.execute(query, values)
        self.conn.commit()

    def update_status(self, item_id: int, status: str):
        """更新状态"""
        self.update_item(item_id, status=status)

    def update_analysis(self, item_id: int, analysis: str):
        """更新分析结果，强制双重标记防止重复处理"""
        self.update_item(
            item_id,
            analysis=analysis,
            status='processed',  # 确保状态变为已处理
            processed_at=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            is_processed=1  # 显式标记为1
        )

    def mark_as_processed(self, item_ids: list):
        """批量标记并防止这些ID再次进入分析流"""
        if not item_ids:
            return
        cursor = self.conn.cursor()
        placeholders = ','.join(['?' for _ in item_ids])
        # 同时更新 status 和 is_processed
        query = f'UPDATE news_items SET is_processed = 1, status = "reported" WHERE id IN ({placeholders})'
        cursor.execute(query, item_ids)
        self.conn.commit()

    def get_stats(self) -> dict:
        """获取统计信息"""
        cursor = self.conn.cursor()

        cursor.execute('SELECT COUNT(*) FROM news_items')
        total = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM news_items WHERE status = "pending"')
        pending = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM news_items WHERE status = "processed"')
        processed = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM news_items WHERE analysis IS NOT NULL')
        analyzed = cursor.fetchone()[0]

        return {
            'total': total,
            'pending': pending,
            'processed': processed,
            'analyzed': analyzed
        }

    # ========================= 自定义关键字管理 =========================

    def get_custom_keywords(self, enabled_only: bool = True) -> List[dict]:
        """获取自定义关键字列表

        Args:
            enabled_only: 是否只返回已启用的关键字

        Returns:
            字典列表，每个字典包含 id, keyword, description, is_enabled, created_at, category
        """
        cursor = self.conn.cursor()
        try:
            if enabled_only:
                cursor.execute('''
                    SELECT id, keyword, description, is_enabled, created_at, category
                    FROM custom_keywords
                    WHERE is_enabled = 1
                    ORDER BY category, created_at DESC
                ''')
            else:
                cursor.execute('''
                    SELECT id, keyword, description, is_enabled, created_at, category
                    FROM custom_keywords
                    ORDER BY category, created_at DESC
                ''')
            rows = cursor.fetchall()

            columns = ['id', 'keyword', 'description', 'is_enabled', 'created_at', 'category']
            return [dict(zip(columns, row)) for row in rows]
        finally:
            cursor.close()

    def add_custom_keyword(self, keyword: str, description: str = '', category: str = 'default') -> int:
        """添加自定义关键字

        Args:
            keyword: 关键字文本
            description: 描述说明
            category: 分类（可选）

        Returns:
            新添加记录的ID，失败返回-1
        """
        from datetime import datetime
        cursor = self.conn.cursor()
        try:
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('''
                INSERT INTO custom_keywords (keyword, description, category, is_enabled, created_at)
                VALUES (?, ?, ?, 1, ?)
            ''', (keyword, description, category, now))
            self.conn.commit()
            return cursor.lastrowid
        except sqlite3.IntegrityError:
            # 关键字已存在
            return -1
        finally:
            cursor.close()

    def delete_custom_keyword(self, keyword_id: int) -> bool:
        """删除自定义关键字

        Args:
            keyword_id: 关键字ID

        Returns:
            是否删除成功
        """
        cursor = self.conn.cursor()
        try:
            cursor.execute('DELETE FROM custom_keywords WHERE id = ?', (keyword_id,))
            self.conn.commit()
            return cursor.rowcount > 0
        finally:
            cursor.close()

    def update_custom_keyword(self, keyword_id: int, **kwargs) -> bool:
        """更新自定义关键字

        Args:
            keyword_id: 关键字ID
            **kwargs: 要更新的字段（keyword, description, is_enabled, category）

        Returns:
            是否更新成功
        """
        if not kwargs:
            return False

        cursor = self.conn.cursor()
        try:
            set_clause = ', '.join([f'{k} = ?' for k in kwargs.keys()])
            query = f'UPDATE custom_keywords SET {set_clause} WHERE id = ?'
            values = list(kwargs.values()) + [keyword_id]
            cursor.execute(query, values)
            self.conn.commit()
            return cursor.rowcount > 0
        finally:
            cursor.close()

    def has_custom_keywords(self) -> bool:
        """检查是否有已启用的自定义关键字

        Returns:
            是否存在已启用的自定义关键字
        """
        cursor = self.conn.cursor()
        try:
            cursor.execute('SELECT COUNT(*) FROM custom_keywords WHERE is_enabled = 1')
            count = cursor.fetchone()[0]
            return count > 0
        finally:
            cursor.close()

    def cleanup_irrelevant_items(self) -> int:
        """清理不匹配自定义关键字的待处理资讯

        将所有 status='pending' 且不匹配当前启用自定义关键字的资讯
        更新为 status='ignored', is_processed=1

        Returns:
            清理的资讯数量
        """
        from datetime import datetime

        # 获取所有已启用的自定义关键字
        custom_keywords = self.get_custom_keywords(enabled_only=True)
        keyword_list = [kw['keyword'] for kw in custom_keywords]

        # 如果没有启用关键字，则不需要清理
        if not keyword_list:
            return 0

        cursor = self.conn.cursor()
        try:
            # 获取所有待处理的资讯
            cursor.execute('SELECT id, title, content FROM news_items WHERE status = "pending"')
            pending_items = cursor.fetchall()

            if not pending_items:
                return 0

            # 找出不匹配关键字的资讯ID
            irrelevant_ids = []
            for item in pending_items:
                item_id, title, content = item
                title = title or ''
                content = content or ''

                # 检查是否包含任一关键字
                matches = False
                for keyword in keyword_list:
                    if keyword.lower() in title.lower() or keyword.lower() in content.lower():
                        matches = True
                        break

                # 如果不匹配任何关键字，则标记为需要清理
                if not matches:
                    irrelevant_ids.append(item_id)

            # 批量更新状态
            if irrelevant_ids:
                now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                placeholders = ','.join(['?' for _ in irrelevant_ids])
                query = f'''
                    UPDATE news_items
                    SET status = 'ignored', is_processed = 1, processed_at = ?
                    WHERE id IN ({placeholders})
                '''
                cursor.execute(query, [now] + irrelevant_ids)
                self.conn.commit()

            return len(irrelevant_ids)
        finally:
            cursor.close()

# ============================================
# 分类引擎
# ============================================

class ClassificationEngine:
    """内容分类引擎"""

    def __init__(self, config_path: str = None, log_callback=None):
        self.config_path = config_path or os.path.join(
            os.path.dirname(__file__), 'classification_rules.yaml'
        )
        self.extra_rules_path = os.path.join(
            os.path.dirname(__file__), 'extra_rules.yaml'
        )
        self.config = None
        self.extra_rules = None
        self.log_callback = log_callback
        self._load_config()
        self._load_extra_rules()

    def _log(self, message: str):
        """日志输出（不打印到控制台）"""
        if self.log_callback:
            self.log_callback(f"[分类] {message}")

    def _load_config(self):
        """加载配置文件"""
        try:
            # 支持PyInstaller打包后的路径
            if getattr(sys, 'frozen', False):
                # 优先使用外部配置
                external_config = os.path.join(
                    os.path.dirname(sys.executable),
                    os.path.basename(self.config_path)
                )
                if os.path.exists(external_config):
                    self.config_path = external_config
                elif '_MEIPASS' in os.environ:
                    # 使用内部配置
                    internal_config = os.path.join(
                        os.environ['_MEIPASS'],
                        os.path.basename(self.config_path)
                    )
                    if os.path.exists(internal_config):
                        self.config_path = internal_config

            if os.path.exists(self.config_path):
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    self.config = yaml.safe_load(f)
                self._log(f"已加载配置: {self.config_path}")
            else:
                self._log(f"配置文件不存在: {self.config_path}")
                self.config = self._get_default_config()
        except Exception as e:
            self._log(f"加载配置失败: {e}")
            self.config = self._get_default_config()

    def _load_extra_rules(self):
        """加载额外规则配置"""
        try:
            # 支持PyInstaller打包后的路径
            config_path = self.extra_rules_path
            if getattr(sys, 'frozen', False):
                # 优先使用外部配置
                external_config = os.path.join(
                    os.path.dirname(sys.executable),
                    'extra_rules.yaml'
                )
                if os.path.exists(external_config):
                    config_path = external_config
                elif '_MEIPASS' in os.environ:
                    # 使用内部配置
                    internal_config = os.path.join(
                        os.environ['_MEIPASS'],
                        'extra_rules.yaml'
                    )
                    if os.path.exists(internal_config):
                        config_path = internal_config

            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    self.extra_rules = yaml.safe_load(f)
                self._log(f"已加载额外规则: {config_path}")
            else:
                self._log(f"额外规则不存在，使用默认值")
                self.extra_rules = self._get_default_extra_rules()
        except Exception as e:
            self._log(f"加载额外规则失败: {e}")
            self.extra_rules = self._get_default_extra_rules()

    def _get_default_extra_rules(self) -> dict:
        """获取默认额外规则"""
        return {
            'intersection_rules': {},
            'dynamic_threshold': {'enabled': False, 'super_sources': []},
            'pre_filter': {'enabled': False, 'advanced_blackbox': []},
            'prompt_injection': {'enabled': False, 'injections': {}},
            'semantic_dedup': {'enabled': False},
            'advanced_settings': {
                'show_score_details': False,
                'show_intersection_logic': False,
                'show_dynamic_threshold': False,
                'show_pre_filter': False
            }
        }

    def _get_default_config(self) -> dict:
        """获取默认配置"""
        return {
            'global_settings': {
                'min_score_threshold': 20,
                'title_weight_multiplier': 2.0,
                'content_weight_multiplier': 1.0,
                'min_keyword_matches': 2
            },
            'blacklists': [
                'promo', 'discount', 'sale', 'coupon',
                'stock price', 'shareholders', 'ipo'
            ],
            'categories': {
                'ai_medical': {
                    'name': 'AI+医疗',
                    'threshold': 20,
                    'keywords': {
                        'AI': 12, '人工智能': 12, 'machine learning': 12,
                        'medical': 12, 'diagnosis': 12, 'healthcare': 12
                    }
                }
            }
        }

    def _parse_config(self) -> dict:
        """解析配置"""
        if not self.config:
            return {}

        categories = self.config.get('categories', {})
        global_settings = self.config.get('global_settings', {})
        blacklists = self.config.get('blacklists', [])
        hard_exclude = self.config.get('hard_exclude', [])

        return {
            'categories': categories,
            'global_settings': global_settings,
            'blacklists': blacklists,
            'hard_exclude': hard_exclude
        }

    def pre_filter_check(self, title: str, content: str = "") -> tuple:
        """预处理过滤检查（Pre-AI Filter）

        Returns:
            (should_skip, skip_reason): (是否跳过, 跳过原因)
        """
        if not self.extra_rules.get('pre_filter', {}).get('enabled', False):
            return False, ""

        advanced_blackbox = self.extra_rules.get('pre_filter', {}).get('advanced_blackbox', [])
        skip_immediately = self.extra_rules.get('pre_filter', {}).get('skip_immediately', [])
        show_details = self.extra_rules.get('advanced_settings', {}).get('show_pre_filter', False)

        text = f"{title} {content}"

        # 检查立即跳过关键词（优先级最高）
        for keyword in skip_immediately:
            if keyword in text:
                if show_details:
                    self._log(f"[预过滤-立即跳过] 命中关键词: {keyword}")
                return True, f"skip_immediately: {keyword}"

        # 检查高级黑名单
        for keyword in advanced_blackbox:
            if keyword in text:
                if show_details:
                    self._log(f"[预过滤-高级黑名单] 命中关键词: {keyword}")
                return True, f"advanced_blackbox: {keyword}"

        return False, ""

    def is_hard_excluded(self, title: str, content: str = "", source: str = "") -> bool:
        """检查是否在硬屏蔽列表中（支持白名单例外，但更保守）"""
        parsed = self._parse_config()
        hard_exclude = parsed.get('hard_exclude', [])

        # 检查标题和内容
        text = f"{title} {content}"

        for keyword in hard_exclude:
            if keyword in text:
                return True

        return False

    def is_blacklisted(self, title: str, content: str, source: str = "") -> bool:
        """检查是否在黑名单中（支持白名单例外）"""
        parsed = self._parse_config()
        blacklists = parsed.get('blacklists', [])
        whitelist = parsed.get('global_settings', {}).get('whitelist', [])

        text = f"{title} {content}".lower()
        source_lower = source.lower()

        # 检查白名单例外
        for whitelist_item in whitelist:
            whitelist_source = whitelist_item.get('source', '').lower()
            if whitelist_source in source_lower:
                # 此来源在白名单中，检查是否允许被过滤的关键词
                allowed_keywords = whitelist_item.get('allowed_keywords', [])
                for keyword in blacklists:
                    if keyword.lower() in text:
                        # 如果关键词在允许列表中，则不过滤
                        if keyword.lower() in [k.lower() for k in allowed_keywords]:
                            return False  # 白名单例外，不过滤
                        # 否则继续检查其他过滤规则
                        return True

        # 非白名单来源，执行正常的黑名单检查
        for keyword in blacklists:
            if keyword.lower() in text:
                return True

        return False

    def get_source_bonus(self, source: str) -> float:
        """获取来源加分"""
        parsed = self._parse_config()
        source_weights = parsed.get('global_settings', {}).get('source_trust_weights', {})

        for src_pattern, weight in source_weights.items():
            if src_pattern.lower() in source.lower():
                return weight

        return 0.0

    def get_dynamic_threshold_adjustment(self, source: str, category_key: str) -> int:
        """获取动态阈值调整（顶级来源降低门槛）"""
        if not self.extra_rules.get('dynamic_threshold', {}).get('enabled', False):
            return 0

        super_sources = self.extra_rules.get('dynamic_threshold', {}).get('super_sources', [])
        max_adjustment = self.extra_rules.get('dynamic_threshold', {}).get('max_adjustment', -8)
        show_details = self.extra_rules.get('advanced_settings', {}).get('show_dynamic_threshold', False)

        adjustment = 0
        for src_config in super_sources:
            src_name = src_config.get('name', '')
            if src_name.lower() in source.lower():
                adj = src_config.get('threshold_adjustment', 0)
                adjustment = min(adjustment, adj)  # 取最小值（最大的负数调整）

        # 限制最大调整幅度
        adjustment = max(adjustment, max_adjustment)

        if adjustment != 0 and show_details:
            self._log(f"[动态阈值] 来源: {source}, 调整: {adjustment}分")

        return adjustment

    def check_intersection_rules(self, category_key: str, title: str, content: str) -> tuple:
        """检查复合命中规则（Intersection Logic）

        Returns:
            (passed, details): (是否通过, 详细信息)
        """
        intersection_rules = self.extra_rules.get('intersection_rules', {})
        category_rule = intersection_rules.get(category_key, {})

        if not category_rule.get('enabled', False):
            return True, "复合命中规则未启用"

        dimensions = category_rule.get('dimensions', {})
        min_scores = category_rule.get('min_scores', {})
        require_all = category_rule.get('require_all', True)
        show_details = self.extra_rules.get('advanced_settings', {}).get('show_intersection_logic', False)

        text = f"{title} {content}".lower()
        dimension_scores = {}

        # 计算每个维度的得分
        for dim_name, keywords in dimensions.items():
            dim_score = 0
            for keyword in keywords:
                if keyword.lower() in text:
                    # 简单的计数，每个关键词算1分
                    dim_score += 1

            dimension_scores[dim_name] = dim_score

        # 检查是否满足要求
        passed = True
        failed_dimensions = []

        for dim_name, min_score in min_scores.items():
            if dimension_scores.get(dim_name, 0) < min_score:
                passed = False
                failed_dimensions.append(f"{dim_name}({dimension_scores.get(dim_name, 0)}/{min_score})")

        if show_details:
            if passed:
                self._log(f"[复合命中] {category_key}: 通过, 维度得分: {dimension_scores}")
            else:
                self._log(f"[复合命中] {category_key}: 未通过, 维度得分: {dimension_scores}, 失败: {failed_dimensions}")

        details = {
            'passed': passed,
            'dimension_scores': dimension_scores,
            'failed_dimensions': failed_dimensions
        }

        return passed, details

    def calculate_score(self, title: str, content: str, source: str) -> dict:
        """计算分类得分（支持动态权重调整）"""
        parsed = self._parse_config()
        categories = parsed.get('categories', {})
        global_settings = parsed.get('global_settings', {})

        title_weight = global_settings.get('title_weight_multiplier', 2.0)
        content_weight = global_settings.get('content_weight_multiplier', 1.0)
        min_threshold = global_settings.get('min_score_threshold', 20)

        results = {}

        for cat_key, cat_config in categories.items():
            keywords = cat_config.get('keywords', {})
            base_threshold = cat_config.get('threshold', min_threshold)

            # 动态阈值调整
            threshold_adj = self.get_dynamic_threshold_adjustment(source, cat_key)
            adjusted_threshold = base_threshold + threshold_adj

            score = 0.0
            text = f"{title} {content}".lower()

            for keyword, weight in keywords.items():
                keyword_lower = keyword.lower()

                # 标题匹配
                if keyword_lower in title.lower():
                    score += weight * title_weight

                # 内容匹配
                count = text.count(keyword_lower)
                if count > 0:
                    score += weight * content_weight * min(count, 3)

            # 来源加分
            score += self.get_source_bonus(source)

            # 检查复合命中规则
            intersection_passed, intersection_details = self.check_intersection_rules(cat_key, title, content)

            results[cat_key] = {
                'score': score,
                'base_threshold': base_threshold,
                'threshold_adjustment': threshold_adj,
                'adjusted_threshold': adjusted_threshold,
                'passed': score >= adjusted_threshold and intersection_passed,
                'intersection_passed': intersection_passed,
                'intersection_details': intersection_details,
                'category_name': cat_config.get('name', cat_key)
            }

        return results

    def classify(self, item: dict) -> dict:
        """分类资讯项"""
        title = item.get('title', '')
        content = item.get('content', '')
        source = item.get('source', '')

        # 检查黑名单（传递 source 参数以支持白名单）
        if self.is_blacklisted(title, content, source):
            return {
                'category': 'blacklisted',
                'category_name': '黑名单',
                'score': 0,
                'passed': False
            }

        # 计算得分
        scores = self.calculate_score(title, content, source)

        # 找到最高分的分类
        best_category = None
        best_score = 0

        for cat_key, result in scores.items():
            if result['passed'] and result['score'] > best_score:
                best_category = cat_key
                best_score = result['score']

        if best_category:
            return {
                'category': best_category,
                'category_name': scores[best_category]['category_name'],
                'score': best_score,
                'passed': True,
                'all_scores': scores
            }
        else:
            return {
                'category': 'other',
                'category_name': '其他',
                'score': best_score,
                'passed': False,
                'all_scores': scores
            }

# ============================================
# RSS采集器
# ============================================

class SimpleCollector:
    """简化的RSS采集器"""

    def __init__(self, db: SimpleDatabase, classifier: ClassificationEngine, log_callback=None):
        self.db = db
        self.classifier = classifier
        self.log_callback = log_callback or print
        self.session = self._create_retry_session()
        self.saved_item_ids = []  # 记录本次采集保存的资讯ID

    def _create_retry_session(self):
        """创建带重试的会话"""
        session = requests.Session()
        adapter = requests.adapters.HTTPAdapter(
            max_retries=urllib3.util.Retry(
                total=3,
                backoff_factor=1,
                status_forcelist=[500, 502, 503, 504]
            )
        )
        session.mount('http://', adapter)
        session.mount('https://', adapter)
        return session

    def _log(self, message: str):
        """日志输出"""
        self.log_callback(f"[采集] {message}")

    def check_source_available(self, url: str) -> bool:
        """检查RSS源是否可用

        注意：此方法在多线程环境中调用，必须为每次调用创建独立的session
        """
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'application/rss+xml, application/xml, text/xml, */*',
            }

            # 为每次调用创建独立的session，避免多线程共享导致连接混乱
            session = requests.Session()
            adapter = requests.adapters.HTTPAdapter(
                max_retries=urllib3.util.Retry(
                    total=3,
                    backoff_factor=1,
                    status_forcelist=[500, 502, 503, 504]
                )
            )
            session.mount('http://', adapter)
            session.mount('https://', adapter)

            response = session.head(url, headers=headers, timeout=10, verify=False)
            if response.status_code in [200, 301, 302]:
                return True

            # 如果HEAD不支持，尝试GET
            response = session.get(url, headers=headers, timeout=10, verify=False, stream=True)
            if response.status_code == 200:
                # 检查是否是有效的RSS/XML
                content_type = response.headers.get('Content-Type', '').lower()
                if 'xml' in content_type or 'rss' in content_type or 'atom' in content_type:
                    return True

                # 读取前100字节检查
                preview = response.raw.read(100)
                if b'<?xml' in preview or b'<rss' in preview or b'<feed' in preview:
                    return True

            return False

        except Exception as e:
            return False

    def get_available_sources(self) -> dict:
        """获取所有可用的RSS源（从数据库读取）"""
        available = {
            'medical': [],
            'ai_tech': [],
            'news': [],
            'movie': [],
            'gaming': []
        }

        # [OK] 资源回收：使用try...finally确保cursor关闭
        cursor = self.db.conn.cursor()
        try:
            cursor.execute('SELECT url, source_type, is_enabled FROM rss_sources WHERE is_enabled = 1')
            sources = cursor.fetchall()
        finally:
            cursor.close()

        if not sources:
            self._log("[WARNING] 数据库中没有RSS源，使用默认源")
            # 如果数据库中没有源，使用硬编码的默认源
            return self._get_default_sources()

        self._log(f"正在检测RSS源可用性...（共{len(sources)}个源）")

        # 使用线程池并发检测
        from concurrent.futures import as_completed

        def test_source(url_info):
            url, source_type = url_info
            if self.check_source_available(url):
                return (source_type, url)
            return None

        # 准备所有源
        all_urls = [(url, source_type) for url, source_type, _ in sources]

        # 并发检测
        with ThreadPoolExecutor(max_workers=20) as executor:
            futures = {executor.submit(test_source, url_info): url_info for url_info in all_urls}

            for future in as_completed(futures):
                try:
                    result = future.result(timeout=15)
                    if result:
                        source_type, url = result
                        available[source_type].append(url)
                except:
                    pass

        # 统计
        total_available = (
            len(available['medical']) +
            len(available['ai_tech']) +
            len(available['news']) +
            len(available['movie']) +
            len(available['gaming'])
        )

        self._log(f"[OK] 可用RSS源: {total_available}/{len(all_urls)}")
        self._log(f"  - 医疗源: {len(available['medical'])} 个")
        self._log(f"  - AI技术源: {len(available['ai_tech'])} 个")
        self._log(f"  - 新闻源: {len(available['news'])} 个")
        self._log(f"  - 影视娱乐源: {len(available['movie'])} 个")
        self._log(f"  - 游戏源: {len(available['gaming'])} 个")

        return available

    def _get_default_sources(self) -> dict:
        """获取硬编码的默认RSS源（后备方案）"""
        all_urls = []
        for url in MEDICAL_RSS_SOURCES:
            all_urls.append((url, 'medical'))
        for url in AI_TECH_RSS_SOURCES:
            all_urls.append((url, 'ai_tech'))
        for url in BREAKING_NEWS_RSS_SOURCES:
            all_urls.append((url, 'news'))
        for url in MOVIE_ENTERTAINMENT_RSS_SOURCES:
            all_urls.append((url, 'movie'))
        for url in GAMING_RSS_SOURCES:
            all_urls.append((url, 'gaming'))

        available = {
            'medical': [],
            'ai_tech': [],
            'news': [],
            'movie': [],
            'gaming': []
        }

        for url, source_type in all_urls:
            if self.check_source_available(url):
                available[source_type].append(url)

        return available

    # ========================= RSS安全解析方法 =========================

    def _sanitize_content(self, content: str) -> str:
        """清理RSS内容：移除BOM头、无效字符、控制字符

        Args:
            content: 原始文本内容

        Returns:
            清理后的文本内容
        """
        if not content:
            return content

        # 移除BOM头（UTF-8 BOM 和 UTF-16 BOM）
        content = content.replace('', '').replace('', '')

        # 移除其他控制字符（保留换行、制表符）
        import re
        # 保留 \t \n \r，移除其他控制字符（0x00-0x1f, 0x7f-0x9f）
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', content)

        return content

    def _validate_rss_content(self, content: str) -> tuple:
        """验证内容是否为有效的RSS/XML

        Args:
            content: 待验证的内容

        Returns:
            (is_valid, error_message)
        """
        if not content or len(content.strip()) < 50:
            return False, "内容为空或过短"

        content_lower = content.lower()

        # 检查是否包含RSS/XML标记
        rss_indicators = ['<rss', '<feed', '<entry', '<?xml', '<rdf:rdf']
        if not any(indicator in content_lower for indicator in rss_indicators):
            return False, "非RSS/XML格式"

        return True, ""

    def _safe_parse_rss(self, url: str, response) -> object:
        """安全解析RSS内容，支持多层降级策略

        降级策略：
        1. 主策略：使用 response.text（已解码的字符串）
        2. 降级1：清理BOM和特殊字符后重试
        3. 降级2：使用 response.content（原始字节流，让feedparser自动检测编码）
        4. 降级3：让feedparser直接从URL拉取（绕过requests）
        5. 兜底：返回None，跳过该源

        Args:
            url: RSS源URL（用于日志和降级策略3）
            response: requests.Response对象

        Returns:
            feedparser.FeedParserDict 对象，或 None（所有策略都失败）
        """
        # 记录尝试策略
        attempts = []

        # ============ 策略1：使用 response.text（推荐） ============
        try:
            content_type = response.headers.get('Content-Type', '')
            text = response.text

            # 预检查
            is_valid, error_msg = self._validate_rss_content(text)
            if is_valid:
                feed = feedparser.parse(text)
                if feed and feed.get('entries', None) is not None:
                    attempts.append("策略1(response.text): 成功")
                    return feed
                else:
                    attempts.append("策略1(response.text): 解析结果无效")
            else:
                    attempts.append(f"策略1(response.text): {error_msg}")
        except Exception as e:
            attempts.append(f"策略1(response.text): {type(e).__name__}: {str(e)[:50]}")

        # ============ 策略2：清理内容后重试 ============
        try:
            text = self._sanitize_content(response.text)
            is_valid, error_msg = self._validate_rss_content(text)
            if is_valid:
                feed = feedparser.parse(text)
                if feed and feed.get('entries', None) is not None:
                    attempts.append("策略2(清理后): 成功")
                    self._log(f"[WARNING] {url} 使用策略2解析成功")
                    return feed
                else:
                    attempts.append(f"策略2(清理后): {error_msg}")
        except Exception as e:
            attempts.append(f"策略2(清理后): {type(e).__name__}: {str(e)[:50]}")

        # ============ 策略3：使用 response.content（原始字节流） ============
        try:
            feed = feedparser.parse(response.content)
            if feed and feed.get('entries', None) is not None:
                attempts.append("策略3(response.content): 成功")
                self._log(f"[WARNING] {url} 使用策略3解析成功")
                return feed
            else:
                attempts.append("策略3(response.content): 解析结果无效")
        except Exception as e:
            attempts.append(f"策略3(response.content): {type(e).__name__}: {str(e)[:50]}")

        # ============ 策略4：让feedparser直接从URL拉取（最后手段）============
        try:
            # 注意：这个策略无法使用自定义headers，但可以绕过requests的某些问题
            feed = feedparser.parse(url)
            if feed and feed.get('entries', None) is not None:
                attempts.append("策略4(直接URL): 成功")
                self._log(f"[WARNING] {url} 使用策略4解析成功（直接从URL拉取）")
                return feed
            else:
                attempts.append("策略4(直接URL): 解析结果无效")
        except Exception as e:
            attempts.append(f"策略4(直接URL): {type(e).__name__}: {str(e)[:50]}")

        # ============ 所有策略都失败 ============
        self._log(f"[ERROR] {url} 所有解析策略均失败:")
        for i, attempt in enumerate(attempts, 1):
            self._log(f"    尝试{i}: {attempt}")

        return None

    def fetch_source(self, url: str, category: str = '未分类') -> dict:
        """采集单个RSS源，返回详细统计（增强版：多层容错）

        改进点：
        1. 响应预检查（状态码、Content-Type、内容长度）
        2. 使用安全解析函数（多层降级策略）
        3. 异常分级（WARNING/ERROR）
        4. 确保单个源失败不影响其他源
        """
        stats = {
            'fetched': 0,      # 采集的总条数
            'duplicates': 0,   # 重复的条数
            'blacklisted': 0,  # 黑名单过滤的条数
            'saved': 0         # 最终保存的条数
        }

        # ========== 第一层：网络请求与预检查 ==========
        try:
            # 设置headers
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,application/rss+xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'gzip, deflate',  # 明确指定支持压缩
            }

            # 发起请求
            response = self.session.get(url, headers=headers, timeout=30, verify=False)
            response.raise_for_status()

            # ========== 预检查1：Content-Type ==========
            content_type = response.headers.get('Content-Type', '').lower()
            if content_type and not any(ct in content_type for ct in ['xml', 'rss', 'atom', 'text/html', 'text/plain']):
                self._log(f"[WARNING] {url} 异常Content-Type: {content_type}")
                # 不返回，继续尝试

            # ========== 预检查2：内容长度 ==========
            content_length = len(response.content)
            if content_length < 100:
                self._log(f"[WARNING] {url} 响应内容过短: {content_length}字节")
                return stats
            if content_length > 50 * 1024 * 1024:  # 50MB
                self._log(f"[WARNING] {url} 响应内容过大: {content_length/1024/1024:.1f}MB")
                return stats

            # ========== 第二层：安全解析RSS ==========
            feed = self._safe_parse_rss(url, response)

            if feed is None:
                # 解析完全失败，已经在 _safe_parse_rss 中记录了错误
                return stats

            # ========== 第三层：验证解析结果 ==========
            if not hasattr(feed, 'entries') or not feed.entries:
                self._log(f"[WARNING] {url} 解析成功但无entries，可能是空RSS源")
                # 尝试记录feed信息以便调试
                if hasattr(feed, 'feed'):
                    feed_title = feed.feed.get('title', '未知')
                    self._log(f"    Feed标题: {feed_title}")
                return stats

            # ========== 第四层：处理条目 ==========
            feed_title = feed.feed.get('title', url)

            for entry in feed.entries:
                try:
                    stats['fetched'] += 1

                    # 检查日期
                    if not self._within_days(entry, 7):
                        continue

                    # 提取内容
                    title = entry.get('title', '')
                    content = self._clean_html(entry.get('description', ''))
                    source = feed_title  # 使用预先提取的feed标题

                    # 基础验证
                    if not title or not title.strip():
                        continue

                    # 预处理过滤
                    should_skip, skip_reason = self.classifier.pre_filter_check(title, content)
                    if should_skip:
                        stats['blacklisted'] += 1
                        if self.classifier.extra_rules.get('advanced_settings', {}).get('show_pre_filter', False):
                            self._log(f"[预过滤跳过] {title[:50]}... 原因: {skip_reason}")
                        continue

                    # 黑名单检查
                    if self.classifier.is_blacklisted(title, content, source):
                        stats['blacklisted'] += 1
                        continue

                    # 硬屏蔽词检查
                    if self.classifier.is_hard_excluded(title, content, source):
                        stats['blacklisted'] += 1
                        continue

                    # 内容价值评估
                    if not self.is_valuable_content(title, content):
                        continue

                    item = {
                        'url': entry.get('link', ''),
                        'title': title,
                        'content': content,
                        'source': source,
                        'pub_date': self._extract_date(entry),
                        'category': category
                    }

                    # URL去重（移除跟踪参数）
                    clean_url = item['url'].split('?')[0].split('#')[0]
                    if self.db.is_duplicate(clean_url):
                        stats['duplicates'] += 1
                        continue

                    # 标题去重（混合相似度算法）
                    if self.db.is_duplicate_by_title(item['title'], threshold=0.85):
                        stats['duplicates'] += 1
                        self._log(f"[拦截重复标题] {item['title'][:30]}")
                        continue

                    # 内容去重（可选）
                    dedup_settings = self.classifier.config.get('deduplication_settings', {})
                    if dedup_settings.get('enable_content_dedup', False):
                        threshold = dedup_settings.get('content_similarity_threshold', 0.85)
                        if self.db.is_duplicate_by_content(item['content'], threshold):
                            stats['duplicates'] += 1
                            self._log(f"[内容去重] {item['title'][:50]}...")
                            continue

                    # 保存资讯
                    item_id = self.db.save_item(item)
                    if item_id > 0:
                        stats['saved'] += 1
                        self.saved_item_ids.append(item_id)

                except Exception as e:
                    # 单个条目处理失败，不影响其他条目
                    continue

            # ========== 第五层：结果汇总 ==========
            if stats['fetched'] > 0:
                self._log(
                    f"[OK] {feed_title}: "
                    f"采集{stats['fetched']}条, "
                    f"重复{stats['duplicates']}条, "
                    f"黑名单{stats['blacklisted']}条, "
                    f"保存{stats['saved']}条"
                )
            else:
                self._log(f"[WARNING] {feed_title} 无可采集内容（可能是日期范围外或格式问题）")

        # ========== 异常分级处理 ==========
        except requests.exceptions.Timeout:
            self._log(f"[ERROR] {url} 请求超时（>30秒）")
        except requests.exceptions.ConnectionError:
            self._log(f"[ERROR] {url} 网络连接失败")
        except requests.exceptions.HTTPError as e:
            self._log(f"[ERROR] {url} HTTP错误: {e.response.status_code}")
        except requests.exceptions.RequestException as e:
            self._log(f"[ERROR] {url} 请求异常: {type(e).__name__}: {str(e)[:100]}")
        except Exception as e:
            # 未知异常，记录详细信息但不中断程序
            self._log(f"[ERROR] {url} 未知错误: {type(e).__name__}: {str(e)[:100]}")

        return stats

    def _within_days(self, entry, days: int = 7) -> bool:
        """检查是否在指定天数内"""
        try:
            pub_date = self._extract_date(entry)
            if pub_date:
                date_obj = datetime.strptime(pub_date, '%Y-%m-%d')
                return (datetime.now() - date_obj).days <= days
            return True
        except:
            return True

    def _clean_html(self, html: str) -> str:
        """清理HTML标签（轻量级，用于一般HTML清理）"""
        if not html:
            return ''

        # 移除HTML标签
        text = re.sub(r'<[^>]+>', ' ', html)
        # 移除多余空格
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _clean_html_pro(self, html_str: str) -> str:
        """更彻底的内容清洗（用于去重前的预处理）

        改进点：
        1. 转换HTML实体字符（&nbsp; &amp; 等）
        2. 移除脚本和样式标签内容
        3. 移除所有HTML标签
        4. 合并多余空格

        Args:
            html_str: 包含HTML的字符串

        Returns:
            清洗后的纯文本
        """
        import html as html_module
        if not html_str:
            return ""

        # 1. 转换实体字符如 &nbsp; &amp; 等
        text = html_module.unescape(html_str)

        # 2. 移除脚本和样式标签内容
        text = re.sub(r'<(script|style).*?>.*?</\1>', '', text, flags=re.DOTALL | re.IGNORECASE)

        # 3. 移除所有HTML标签
        text = re.sub(r'<[^>]+>', ' ', text)

        # 4. 合并多余空格
        return re.sub(r'\s+', ' ', text).strip()

    def _extract_date(self, entry) -> str:
        """提取发布日期"""
        try:
            if hasattr(entry, 'published_parsed') and entry.published_parsed:
                return datetime(*entry.published_parsed[:6]).strftime('%Y-%m-%d')
            elif hasattr(entry, 'updated_parsed') and entry.updated_parsed:
                return datetime(*entry.updated_parsed[:6]).strftime('%Y-%m-%d')
            return datetime.now().strftime('%Y-%m-%d')
        except:
            return datetime.now().strftime('%Y-%m-%d')

    def is_valuable_content(self, title: str, content: str) -> bool:
        """检查是否为有价值内容"""
        # 检查标题长度
        if len(title) < 10:
            return False

        # 【修改】降低内容长度要求：50 → 20（适应游戏/影视资讯的简短描述）
        if len(content) < 20:
            return False

        # 注意：黑名单检查已在调用此方法之前完成，此处无需重复检查

        return True

    def fetch_all(self) -> dict:
        """采集所有RSS源（只采集可用的源），返回详细统计"""
        # 清空之前保存的ID列表
        self.saved_item_ids = []

        # 汇总统计
        total_stats = {
            'fetched': 0,
            'duplicates': 0,
            'blacklisted': 0,
            'saved': 0,
            'item_ids': []
        }

        # 先检测可用源
        available = self.get_available_sources()

        # 如果没有可用源，返回0
        total_available = (
            len(available['medical']) +
            len(available['ai_tech']) +
            len(available['news']) +
            len(available['movie']) +
            len(available['gaming'])
        )

        if total_available == 0:
            self._log("[WARNING] 没有可用的RSS源！")
            return total_stats

        self._log(f"开始采集 {total_available} 个可用RSS源...")

        # 使用线程池并发采集
        with ThreadPoolExecutor(max_workers=10) as executor:
            futures = {}

            # 只采集可用的医疗源
            for url in available['medical']:
                future = executor.submit(self.fetch_source, url, '医疗')
                futures[future] = ('医疗', url)

            # 只采集可用的AI技术源
            for url in available['ai_tech']:
                future = executor.submit(self.fetch_source, url, 'AI技术')
                futures[future] = ('AI技术', url)

            # 只采集可用的新闻源
            for url in available['news']:
                future = executor.submit(self.fetch_source, url, '新闻')
                futures[future] = ('新闻', url)

            # 【新增】只采集可用的影视娱乐源
            for url in available['movie']:
                future = executor.submit(self.fetch_source, url, '影视娱乐')
                futures[future] = ('影视娱乐', url)

            # 【新增】只采集可用的游戏源
            for url in available['gaming']:
                future = executor.submit(self.fetch_source, url, '游戏')
                futures[future] = ('游戏', url)

            # 等待完成并汇总统计
            for future in as_completed(futures):
                try:
                    stats = future.result()
                    total_stats['fetched'] += stats['fetched']
                    total_stats['duplicates'] += stats['duplicates']
                    total_stats['blacklisted'] += stats['blacklisted']
                    total_stats['saved'] += stats['saved']
                except Exception as e:
                    self._log(f"采集失败: {e}")

        # 显示汇总统计
        self._log("")
        self._log("=" * 60)
        self._log("【采集统计汇总】")
        self._log("=" * 60)
        self._log(f" 总采集条数: {total_stats['fetched']} 条")
        self._log(f" 重复过滤: {total_stats['duplicates']} 条")
        self._log(f" 黑名单过滤: {total_stats['blacklisted']} 条")
        self._log(f"[SAVE] 最终保存: {total_stats['saved']} 条")
        self._log("=" * 60)

        # 计算过滤率
        if total_stats['fetched'] > 0:
            filter_rate = ((total_stats['duplicates'] + total_stats['blacklisted']) / total_stats['fetched']) * 100
            self._log(f"[REPORT] 过滤率: {filter_rate:.1f}%")

        self._log(f"[OK] 采集完成！")

        # 保存本次采集的ID列表
        total_stats['item_ids'] = self.saved_item_ids.copy()

        return total_stats

# ============================================
# 内容处理器
# ============================================

class SimpleProcessor:
    """简化的内容处理器"""

    def __init__(self, db: SimpleDatabase, log_callback=None):
        self.db = db
        self.log_callback = log_callback or print

        # 加载分类配置（包含hard_exclude）
        self.classification_config = self._load_classification_config()

        # 加载提示词配置（在_log方法定义之后调用）
        self.prompts_config = self._load_prompts_config()

        # 加载额外规则配置（包含动态注入规则）
        self.extra_rules = self._load_extra_rules()

        # API配置（从prompts_config.yaml读取，如果配置文件中没有则使用默认值）
        api_config = self.prompts_config.get('api', {})
        self.api_key = api_config.get('api_key', '')  # 移除硬编码密钥，默认为空
        self.api_url = api_config.get('api_url', 'https://open.bigmodel.cn/api/paas/v4/')
        self.model = api_config.get('model', 'glm-4-flash')

        # 记录API配置信息（不记录完整密钥，只记录前8位和后4位）
        api_key_masked = self.api_key[:8] + '...' + self.api_key[-4:] if len(self.api_key) > 12 else '***'
        self.log_callback(f"API配置: URL={self.api_url}, Model={self.model}, Key={api_key_masked}")

    def _log(self, message: str):
        """日志输出"""
        try:
            self.log_callback(f"[处理] {message}")
        except:
            # 如果回调失败，直接打印
            print(f"[处理] {message}")

    def _load_classification_config(self):
        """加载分类配置文件（包含hard_exclude）"""
        try:
            # 获取基础目录
            if getattr(sys, 'frozen', False):
                # PyInstaller打包后的exe
                exe_dir = os.path.dirname(sys.executable)

                # 优先查找外部配置文件（exe所在目录）
                external_config = os.path.join(exe_dir, 'classification_rules.yaml')
                if os.path.exists(external_config):
                    self._log(f"使用外部分类配置: {external_config}")
                    with open(external_config, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)

                # 如果外部没有，查找内部打包的配置文件
                if '_MEIPASS' in os.environ:
                    internal_config = os.path.join(os.environ['_MEIPASS'], 'classification_rules.yaml')
                    if os.path.exists(internal_config):
                        self._log(f"使用内部分类配置: {internal_config}")
                        with open(internal_config, 'r', encoding='utf-8') as f:
                            return yaml.safe_load(f)

                # 都没有，返回空配置
                self._log("未找到分类配置，使用默认值")
                return self._get_default_classification_config()
            else:
                # 开发模式：从脚本目录加载
                config_path = os.path.join(
                    os.path.dirname(__file__), 'classification_rules.yaml'
                )
                if os.path.exists(config_path):
                    self._log(f"使用分类配置: {config_path}")
                    with open(config_path, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)
                else:
                    self._log(f"分类配置不存在: {config_path}，使用默认值")
                    return self._get_default_classification_config()
        except Exception as e:
            self._log(f"加载分类配置失败: {e}，使用默认值")
            return self._get_default_classification_config()

    def _get_default_classification_config(self) -> dict:
        """获取默认分类配置"""
        return {
            'global_settings': {},
            'hard_exclude': [],
            'deduplication_settings': {
                'similarity_threshold': 0.90,
                'enable_title_dedup': True
            }
        }

    def _load_extra_rules(self):
        """加载额外规则配置"""
        try:
            # 获取基础目录
            if getattr(sys, 'frozen', False):
                # PyInstaller打包后的exe
                exe_dir = os.path.dirname(sys.executable)

                # 优先查找外部配置文件（exe所在目录）
                external_config = os.path.join(exe_dir, 'extra_rules.yaml')
                if os.path.exists(external_config):
                    self._log(f"使用外部额外规则: {external_config}")
                    with open(external_config, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)

                # 如果外部没有，查找内部打包的配置文件
                if '_MEIPASS' in os.environ:
                    internal_config = os.path.join(os.environ['_MEIPASS'], 'extra_rules.yaml')
                    if os.path.exists(internal_config):
                        self._log(f"使用内部额外规则: {internal_config}")
                        with open(internal_config, 'r', encoding='utf-8') as f:
                            return yaml.safe_load(f)

                # 都没有，返回空配置
                self._log("未找到额外规则配置，使用默认值")
                return self._get_default_extra_rules()
            else:
                # 开发模式：从脚本目录加载
                config_path = os.path.join(
                    os.path.dirname(__file__), 'extra_rules.yaml'
                )
                if os.path.exists(config_path):
                    self._log(f"使用额外规则: {config_path}")
                    with open(config_path, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)
                else:
                    self._log(f"额外规则不存在: {config_path}，使用默认值")
                    return self._get_default_extra_rules()
        except Exception as e:
            self._log(f"加载额外规则失败: {e}，使用默认值")
            return self._get_default_extra_rules()

    def _get_default_extra_rules(self) -> dict:
        """获取默认额外规则"""
        return {
            'prompt_injection': {
                'enabled': False,
                'dynamic_injections': [],
                'injections': {}
            }
        }

    def _load_prompts_config(self):
        """加载提示词配置文件"""
        try:
            # 获取基础目录
            if getattr(sys, 'frozen', False):
                # PyInstaller打包后的exe
                exe_dir = os.path.dirname(sys.executable)

                # 优先查找外部配置文件（exe所在目录）
                external_config = os.path.join(exe_dir, 'prompts_config.yaml')
                if os.path.exists(external_config):
                    self._log(f"使用外部配置文件: {external_config}")
                    with open(external_config, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)

                # 如果外部没有，查找内部打包的配置文件
                if '_MEIPASS' in os.environ:
                    internal_config = os.path.join(os.environ['_MEIPASS'], 'prompts_config.yaml')
                    if os.path.exists(internal_config):
                        self._log(f"使用内部配置文件: {internal_config}")
                        with open(internal_config, 'r', encoding='utf-8') as f:
                            config = yaml.safe_load(f)

                        # 同时生成外部配置文件模板
                        self._generate_config_template(external_config, config)
                        return config

                # 都没有，生成外部配置文件模板
                self._log(f"未找到配置文件，自动生成模板: {external_config}")
                default_config = self._get_default_prompts_config()
                self._generate_config_template(external_config, default_config)
                return default_config
            else:
                # 开发模式：从脚本目录加载
                config_path = os.path.join(
                    os.path.dirname(__file__), 'prompts_config.yaml'
                )
                if os.path.exists(config_path):
                    self._log(f"使用配置文件: {config_path}")
                    with open(config_path, 'r', encoding='utf-8') as f:
                        return yaml.safe_load(f)
                else:
                    self._log(f"配置文件不存在: {config_path}，自动生成模板")
                    default_config = self._get_default_prompts_config()
                    self._generate_config_template(config_path, default_config)
                    return default_config
        except Exception as e:
            self._log(f"加载提示词配置失败: {e}，使用默认配置")
            return self._get_default_prompts_config()

    def _generate_config_template(self, config_path: str, config: dict):
        """生成配置文件模板"""
        try:
            # 确保config中包含完整的api配置
            if 'api' not in config:
                config['api'] = {
                    'api_url': 'https://open.bigmodel.cn/api/paas/v4/',
                    'api_key': '',
                    'model': 'glm-4-flash',
                    'translation_temperature': 0.3,
                    'analysis_temperature': 0.7,
                    'translation_timeout': 30,
                    'analysis_timeout': 60
                }

            # 写入配置文件
            with open(config_path, 'w', encoding='utf-8') as f:
                yaml.safe_dump(config, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

            self._log(f"[OK] 已生成配置文件模板: {config_path}")
        except Exception as e:
            self._log(f" 生成配置文件模板失败: {e}")

    def _get_default_prompts_config(self) -> dict:
        """获取默认提示词配置"""
        return {
            'translation': {
                'en_to_zh_prompt': '请将以下英文内容翻译成中文，保持专业术语的准确性：\n\n{text}',
                'enable_auto_skip': True
            },
            'analysis': {
                'prompt_template': '''
请分析以下资讯：

标题: {title}

内容: {content}

【第一步：内容审查】
在进行深度分析前，请先执行内容价值评估：

判定准则：如果资讯内容属于"产品上市、销售、促销、预约、抢购、价格信息（如：某手机上架、开启预约、首销、起售价、促销降价）"，且不包含技术创新、功能更新或行业趋势。

执行动作：请直接回复"[跳过：上市促销类资讯]"，不要生成任何分析内容。

注意：技术更新、新功能发布、系统升级等内容应该保留并进行深度分析。

如果资讯内容有价值，请继续第二步。

【第二步：深度分析】
请从以下几个维度进行分析：
1. 技术创新维度
2. 产业影响维度
3. 政策支持维度
4. 经济影响维度
5. 外部环境维度

每个维度给出0-3分的评分，并说明理由。
''',
                'max_content_length': 1000,
                'skip_keywords': ['[跳过', '上市促销']
            },
            'api': {
                'api_url': 'https://open.bigmodel.cn/api/paas/v4/',
                'api_key': '',
                'model': 'glm-4-flash',
                'translation_temperature': 0.3,
                'analysis_temperature': 0.7,
                'translation_timeout': 30,
                'analysis_timeout': 60
            }
        }

    def should_skip_by_keywords(self, title: str) -> bool:
        """根据配置文件中的hard_exclude关键词判断是否应该跳过"""
        hard_exclude = self.classification_config.get('hard_exclude', [])
        for word in hard_exclude:
            if word in title:
                return True
        return False

    def is_valuable_content(self, title: str, content: str) -> bool:
        """检查是否为有价值内容"""
        # 简单检查
        return len(title) > 10 and len(content) > 50

    def process_all(self, limit: int = None, item_ids: list = None) -> int:
        """处理资讯

        Args:
            limit: 限制处理的数量
            item_ids: 只处理指定的资讯ID列表（如果提供，则忽略limit）
        """
        # [OK] 单游标复用：在所有循环外创建一个cursor，整个方法共享
        cursor = self.db.conn.cursor()
        try:
            if item_ids:
                # 只处理指定的ID
                items = []
                # [OK] 复用游标处理所有item_id
                for item_id in item_ids:
                    cursor.execute('SELECT * FROM news_items WHERE id = ? AND status = "pending"', (item_id,))
                    row = cursor.fetchone()
                    if row:
                        columns = [desc[0] for desc in cursor.description]
                        items.append(dict(zip(columns, row)))
                self._log(f"开始处理本次采集的 {len(items)} 条新资讯...")
            else:
                # 处理所有待处理的资讯
                items = self.db.get_raw_items(status='pending', limit=limit)
                self._log(f"开始处理 {len(items)} 条资讯...")

            # 【新增】检查是否设置了自定义关键字
            custom_keywords = self.db.get_custom_keywords(enabled_only=True)
            use_custom_filter = len(custom_keywords) > 0

            if use_custom_filter:
                keyword_list = [kw['keyword'] for kw in custom_keywords]
                self._log(f"[自定义关键字过滤] 已启用，共 {len(keyword_list)} 个关键字")
                self._log(f"  关键字列表: {', '.join(keyword_list[:10])}{'...' if len(keyword_list) > 10 else ''}")

                # 过滤：只保留标题包含任一自定义关键字的资讯
                filtered_items = []
                for item in items:
                    title = item.get('title', '')
                    content = item.get('content', '')
                    # 检查标题和内容是否包含任一关键字
                    for keyword in keyword_list:
                        if keyword.lower() in title.lower() or keyword.lower() in content.lower():
                            filtered_items.append(item)
                            break

                skipped_by_filter = len(items) - len(filtered_items)
                if skipped_by_filter > 0:
                    self._log(f"[自定义关键字过滤] 已跳过 {skipped_by_filter} 条不相关的资讯")

                items = filtered_items
                self._log(f"[自定义关键字过滤] 筛选后剩余 {len(items)} 条资讯待处理")
            else:
                self._log(f"[自定义关键字过滤] 未启用，使用默认过滤逻辑")

            processed_count = 0
            skipped_count = 0

            for item in items:
                try:
                    # 获取原文
                    original_title = item.get('title', '')
                    original_content = item.get('content', '')

                    # 【新增】Token 止损开关：内容长度极端检测
                    # 很多重复内容其实是抓取到了"404"或"订阅提示"页面，这会极大消耗 Token
                    if len(original_content) < 100 and "subscribe" in original_content.lower():
                        self._log(f"[垃圾内容拦截] 可能是订阅墙，跳过处理")
                        self.db.update_item(item['id'], status='skipped', is_processed=1)
                        continue

                    # 【新增】关键词预过滤
                    if self.should_skip_by_keywords(original_title):
                        self._log(f"[跳过] 识别为上市促销类信息: {original_title[:50]}...")
                        skipped_count += 1
                        # 标记为已处理但不保存分析
                        # [OK] 复用游标
                        cursor.execute('UPDATE news_items SET status = "processed" WHERE id = ?', (item['id'],))
                        self.db.conn.commit()
                        continue

                    # 翻译
                    self._log(f"正在翻译: {original_title[:50]}...")
                    translated = self._translate(original_title, original_content)

                    # 分析（使用翻译后的内容）
                    self._log(f"正在分析: {translated['title'][:50]}...")
                    # 获取分类信息
                    category_key = item.get('category_key', item.get('category', ''))
                    analysis = self._deep_analyze(translated['title'], translated['content'], category_key)

                    # 【新增】检查AI是否判定为跳过
                    if analysis is None:
                        self._log(f"[跳过] AI判定为上市促销类资讯: {translated['title'][:50]}...")
                        skipped_count += 1
                        # 标记为已处理但不保存分析
                        # [OK] 复用游标
                        cursor.execute('UPDATE news_items SET status = "processed" WHERE id = ?', (item['id'],))
                        self.db.conn.commit()
                        continue

                    # 生成总结（使用翻译后的内容）
                    self._log(f"正在生成总结: {translated['title'][:50]}...")
                    summary = self._generate_summary(translated['title'], translated['content'])

                    # 保存结果到news_items表
                    # 保存原文到 title_original 和 content_original
                    # 保存译文到 title 和 content
                    # 保存分析到 analysis
                    # 保存总结到 summary
                    # [OK] 复用游标
                    cursor.execute('''
                        UPDATE news_items
                        SET title_original = ?,
                            content_original = ?,
                            title = ?,
                            content = ?,
                            analysis = ?,
                            summary = ?,
                            status = 'processed',
                            translated_at = ?,
                            processed_at = ?,
                            is_processed = 1
                        WHERE id = ?
                    ''', (
                        original_title,
                        original_content,
                        translated['title'],
                        translated['content'],
                        analysis,
                        summary,
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        item['id']
                    ))
                    self.db.conn.commit()

                    processed_count += 1
                    self._log(f"[OK] 已处理: {translated['title'][:50]}...")

                except Exception as e:
                    self._log(f"处理失败: {e}")
                    continue
        finally:
            # [OK] 统一归还游标资源
            cursor.close()

        self._log(f"处理完成，共处理 {processed_count} 条资讯，跳过 {skipped_count} 条")
        return processed_count

    def _translate(self, title: str, content: str) -> dict:
        """翻译内容（使用GLM-4 API）"""
        # 翻译标题
        title_translated = self._translate_text(title)

        # 翻译内容（如果内容太长，分段翻译）
        content_translated = content
        if content and len(content) > 100:
            # 将内容分段，每段不超过2000字符
            chunks = []
            chunk_size = 2000
            for i in range(0, len(content), chunk_size):
                chunk = content[i:i+chunk_size]
                translated_chunk = self._translate_text(chunk)
                chunks.append(translated_chunk)
            content_translated = ''.join(chunks)
        elif content:
            content_translated = self._translate_text(content)

        return {
            'title': title_translated,
            'content': content_translated
        }

    def _translate_text(self, text: str) -> str:
        """翻译单段文本"""
        if not text or not text.strip():
            return text

        # 检查是否主要是英文
        if self._is_english(text):
            # 使用配置中的提示词模板
            prompt_template = self.prompts_config.get('translation', {}).get('en_to_zh_prompt')
            if prompt_template:
                prompt = prompt_template.format(text=text)
            else:
                # 默认提示词
                prompt = f"请将以下英文内容翻译成中文，保持专业术语的准确性：\n\n{text}"
        else:
            # 如果已经是中文，直接返回
            return text

        try:
            # 获取配置中的超时时间
            timeout = self.prompts_config.get('api', {}).get('translation_timeout', 30)
            temperature = self.prompts_config.get('api', {}).get('translation_temperature', 0.3)

            response = requests.post(
                f"{self.api_url}chat/completions",
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": temperature
                },
                timeout=timeout
            )

            if response.status_code == 200:
                result = response.json()
                translated = result['choices'][0]['message']['content'].strip()
                # 移除可能的引号
                if translated.startswith('"') and translated.endswith('"'):
                    translated = translated[1:-1]
                return translated
            else:
                return text

        except Exception as e:
            self._log(f"翻译失败: {e}")
            return text

    def _is_english(self, text: str) -> bool:
        """检查文本是否主要是英文"""
        if not text:
            return False

        # 统计英文字符和中文字符的比例
        english_chars = sum(1 for c in text if c.isascii() and c.isalpha())
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')

        # 如果英文字符多于中文字符，认为是英文
        return english_chars > chinese_chars

    def _generate_summary(self, title: str, content: str) -> str:
        """生成内容总结（2-3句话概括文章主要内容）"""
        # 限制内容长度，避免token消耗过大
        content_preview = content[:1500] if content else ""

        prompt = f"""请用2-3句话总结以下文章的主要内容，要求简洁明了：

标题: {title}
内容: {content_preview}

要求：
1. 总结文章的核心内容和发现
2. 使用简练的中文表达
3. 不要包含评分、维度分析等内容
4. 直接输出总结，不要有其他说明
"""

        try:
            # 获取配置中的超时时间
            timeout = self.prompts_config.get('api', {}).get('translation_timeout', 30)
            temperature = self.prompts_config.get('api', {}).get('translation_temperature', 0.3)

            response = requests.post(
                f"{self.api_url}chat/completions",
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": temperature
                },
                timeout=timeout
            )

            if response.status_code == 200:
                result = response.json()
                summary = result['choices'][0]['message']['content'].strip()
                return summary
            else:
                return "总结生成失败"

        except Exception as e:
            self._log(f"总结生成失败: {e}")
            return f"总结生成异常: {str(e)}"

    def _deep_analyze(self, title: str, content: str, category_key: str = '') -> str:
        """深度分析内容（支持动态和分类提示词增强注入）"""
        # 使用配置中的提示词模板
        prompt_template = self.prompts_config.get('analysis', {}).get('prompt_template')
        max_length = self.prompts_config.get('analysis', {}).get('max_content_length', 1000)

        # 【优先级1】上下文暗示（Context Hints）- 最高优先级
        enhanced_content = content
        injected_instruction = None

        # 从 extra_rules 获取上下文暗示配置
        context_hints = self.extra_rules.get('context_hints', {})
        if context_hints.get('enabled', False):
            trigger_keywords = context_hints.get('trigger_keywords', [])
            text_to_check = f"{title} {content}".lower()

            # 检查是否命中任何触发关键词
            matched_keywords = [kw for kw in trigger_keywords if kw.lower() in text_to_check]

            if matched_keywords:
                # 命中！使用上下文暗示的系统提示
                injected_instruction = context_hints.get('system_prompt', '')
                if injected_instruction:
                    self._log(f"[上下文暗示] 命中关键词: {matched_keywords}, 强制注入科学提示")

        # 【优先级2】动态提示词注入（基于内容特征）
        # 从 extra_rules 获取动态注入配置
        prompt_injection = self.extra_rules.get('prompt_injection', {})
        dynamic_injections_enabled = prompt_injection.get('enabled', False)
        dynamic_injections = prompt_injection.get('dynamic_injections', [])

        # 只有在没有上下文暗示时才检查动态注入
        if not injected_instruction and dynamic_injections_enabled and dynamic_injections:
            # 检查是否匹配任何动态注入规则
            text_to_check = f"{title} {content}".lower()

            for injection_rule in dynamic_injections:
                if not injection_rule.get('enabled', True):
                    continue

                # 检查触发关键词
                trigger_keywords = injection_rule.get('trigger_keywords', [])
                matched_keywords = [kw for kw in trigger_keywords if kw.lower() in text_to_check]

                if matched_keywords:
                    # 命中！使用这个注入指令
                    injected_instruction = injection_rule.get('instruction', '')
                    self._log(f"[动态注入] 匹配规则: {injection_rule['name']}, 命中关键词: {matched_keywords}")
                    break

        # 【优先级3】基于分类的注入（专家指令）
        # 如果没有匹配上下文暗示和动态注入，检查基于分类的注入
        if not injected_instruction and prompt_injection.get('enabled', False):
            injections = prompt_injection.get('injections', {})
            if category_key and category_key in injections:
                category_injection = injections[category_key]
                injected_instruction = category_injection.get('instruction', '')
                if injected_instruction:
                    self._log(f"[分类注入] 使用分类 {category_key} 的专家指令")

        # 如果有注入指令，注入到内容头部
        if injected_instruction:
            enhanced_content = f"{injected_instruction}\n\n【原文内容】\n标题: {title}\n内容: {content}"

        # 使用增强后的内容（但限制长度）
        content_preview = enhanced_content[:max_length] if enhanced_content else ""

        if prompt_template:
            # 使用配置文件中的提示词
            prompt = prompt_template.format(title=title, content=content_preview)
        else:
            # 使用默认提示词
            prompt = f"""
请分析以下资讯：

标题: {title}

内容: {content_preview}

【第一步：内容审查】
在进行深度分析前，请先执行内容价值评估：

判定准则：如果资讯内容属于"产品上市、销售、促销、预约、抢购、价格信息（如：某手机上架、开启预约、首销、起售价、促销降价）"，且不包含技术创新、功能更新或行业趋势。

执行动作：请直接回复"[跳过：上市促销类资讯]"，不要生成任何分析内容。

注意：技术更新、新功能发布、系统升级等内容应该保留并进行深度分析。

如果资讯内容有价值，请继续第二步。

【第二步：深度分析】
请从以下几个维度进行分析：
1. 技术创新维度
2. 产业影响维度
3. 政策支持维度
4. 经济影响维度
5. 外部环境维度

每个维度给出0-3分的评分，并说明理由。
"""

        try:
            # 获取配置中的参数
            timeout = self.prompts_config.get('api', {}).get('analysis_timeout', 60)
            temperature = self.prompts_config.get('api', {}).get('analysis_temperature', 0.7)
            skip_keywords = self.prompts_config.get('analysis', {}).get('skip_keywords', ['[跳过', '上市促销'])

            response = requests.post(
                f"{self.api_url}chat/completions",
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": temperature
                },
                timeout=timeout
            )

            if response.status_code == 200:
                result = response.json()
                analysis = result['choices'][0]['message']['content']

                # 检查是否被AI判定为跳过（使用配置中的关键词）
                for keyword in skip_keywords:
                    if keyword in analysis:
                        return None  # 返回None表示应该跳过

                return analysis
            else:
                return "分析失败"

        except Exception as e:
            return f"分析异常: {str(e)}"

    def _analyze_content(self, title: str, content: str) -> dict:
        """分析内容维度"""
        return {
            '技术/创新': 0,
            '产业/行业': 0,
            '政策支持': 0,
            '经济影响': 0,
            '外部环境': 0
        }

    def _generate_report(self, items: List[dict]) -> str:
        """生成报告"""
        report = []
        report.append("=" * 60)
        report.append("资讯分析报告")
        report.append("=" * 60)
        report.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"资讯总数: {len(items)}")
        report.append("")

        for item in items:
            # 基本信息
            report.append(f"标题: {item.get('title', '无标题')}")
            report.append(f"来源: {item.get('source', '未知')}")
            report.append(f"日期: {item.get('pub_date', item.get('published_date', '未知'))}")
            report.append(f"链接: {item.get('link', '无')}")

            # 内容
            content = item.get('content', '')
            if content:
                # 限制内容长度
                if len(content) > 500:
                    content = content[:500] + "..."
                report.append(f"内容: {content}")

            # 分析结果
            analysis = item.get('analysis', '')
            if analysis:
                report.append(f"分析: {analysis}")

            report.append("-" * 40)

        return "\n".join(report)

    def export_to_word(self, output_path: str, items: List[dict] = None):
        """导出到Word文档"""
        if items is None:
            items = self.db.get_processed_items()

        doc = Document()

        # 标题
        title = doc.add_heading('资讯深度分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生成时间
        doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d")}')
        doc.add_paragraph(f'总计处理: {len(items)} 条资讯')
        doc.add_paragraph()

        # 添加内容
        for item in items:
            doc.add_heading(item.get('title', '无标题'), level=2)
            doc.add_paragraph(f"来源: {item.get('source', '未知')}")
            doc.add_paragraph(f"日期: {item.get('pub_date', item.get('published_date', '未知'))}")
            doc.add_paragraph(f"分类: {item.get('category_key', item.get('category', '未分类'))}")
            doc.add_paragraph(f"链接: {item.get('link', '无')}")

            if item.get('analysis'):
                doc.add_paragraph("分析结果:")
                doc.add_paragraph(item['analysis'])

            doc.add_paragraph()

        # 保存
        doc.save(output_path)
        self._log(f"报告已保存: {output_path}")

# ============================================
# GUI界面
# ============================================

class NewsAppGUI:
    """资讯采集系统GUI"""

    def __init__(self):
        # 设置全局主题为亮色模式
        ctk.set_appearance_mode("Light")
        ctk.set_default_color_theme("blue")

        self.root = ctk.CTk()
        self.root.title("资讯采集系统 V36")
        self.root.geometry("900x700")

        # 配置主窗口背景色（浅灰色）
        self.root.configure(fg_color="#F2F2F2")

        # ========== 小白友好化：检查并自动创建配置文件 ==========
        self._check_and_create_config_file()

        # 数据库路径（支持PyInstaller打包后使用exe所在目录）
        if getattr(sys, 'frozen', False):
            # PyInstaller打包后的exe：使用exe所在目录
            exe_dir = os.path.dirname(sys.executable)
            self.db_path = os.path.join(exe_dir, 'news_v3.db')
        else:
            # 开发环境：使用脚本所在目录
            self.db_path = os.path.join(
                os.path.dirname(__file__), 'news_v3.db'
            )

        # 日志缓冲（用于保存到文件）
        self.log_buffer = []

        # ========== 小白友好化：数据库自动创建（CREATE TABLE IF NOT EXISTS）==========
        self.log("=" * 60)
        self.log(" 资讯采集系统 V3.6 (独立版) 启动中...")
        self.log("=" * 60)

        # 初始化组件（SimpleDatabase会自动创建数据库表）
        self.db = SimpleDatabase(self.db_path)
        self.classifier = ClassificationEngine(log_callback=self.log)
        self.collector = SimpleCollector(self.db, self.classifier, self.log)
        self.processor = SimpleProcessor(self.db, self.log)

        # 创建界面
        self._create_widgets()

        # 输出格式
        self.output_format = "txt"

    def _create_widgets(self):
        """创建界面组件（纯静态左右分栏布局）"""
        # ========== 第1步：创建左侧导航栏 ==========
        sidebar_frame = ctk.CTkFrame(
            self.root,
            fg_color="#E0E0E0",  # 深一点的浅灰色
            corner_radius=0,
            width=180  # 固定宽度180px，足够显示完整文字
        )
        sidebar_frame.pack(side="left", fill="y")
        sidebar_frame.pack_propagate(False)

        # 按钮样式配置
        button_style = {
            "fg_color": "#1F6AA5",  # 专业蓝
            "hover_color": "#3B8ED0",
            "text_color": "#FFFFFF",
            "corner_radius": 8,
            "font": ctk.CTkFont(size=13, family="Microsoft YaHei"),
            "height": 40,
            "anchor": "center"  # 文字居中对齐
        }

        # 前8个功能按钮（垂直排列）
        buttons = [
            ("[全] 全流程运行", self.run_full),
            ("[采] 仅采集", self.run_collect),
            ("[处] 仅处理", self.run_process),
            ("[报] 生成报告", self.run_report),
            ("[源] RSS源", self.open_rss_manager),
            ("[关] 关键字", self.open_keyword_manager),
            ("[提] 提取源代码", self.extract_source_code),
            ("[清] 清理不相关资讯", self.cleanup_irrelevant),
        ]

        for text, command in buttons:
            btn = ctk.CTkButton(
                sidebar_frame,
                text=text,
                command=command,
                **button_style
            )
            btn.pack(side="top", fill="x", padx=10, pady=5)

        # 中间透明占位组件，把设置按钮推到底部
        spacer = ctk.CTkFrame(sidebar_frame, fg_color="transparent")
        spacer.pack(side="top", fill="both", expand=True)

        # 设置按钮（沉底显示）
        settings_btn = ctk.CTkButton(
            sidebar_frame,
            text="[设] 设置",
            command=self.open_settings,
            **button_style
        )
        settings_btn.pack(side="bottom", fill="x", padx=10, pady=15)

        # ========== 第2步：创建右侧内容区 ==========
        content_frame = ctk.CTkFrame(self.root, fg_color="transparent")
        content_frame.pack(side="right", fill="both", expand=True)

        # 主容器 - 白色背景，圆角
        main_container = ctk.CTkFrame(
            content_frame,
            fg_color="#FFFFFF",
            corner_radius=10
        )
        main_container.pack(fill="both", expand=True, padx=15, pady=15)

        # 标题
        title_label = ctk.CTkLabel(
            main_container,
            text="资讯采集系统 V36",
            font=ctk.CTkFont(size=28, weight="bold", family="Microsoft YaHei"),
            text_color="#333333"
        )
        title_label.pack(pady=(20, 15))

        # 副标题
        subtitle_label = ctk.CTkLabel(
            main_container,
            text="智能资讯采集与分析系统",
            font=ctk.CTkFont(size=14, family="Microsoft YaHei"),
            text_color="#666666"
        )
        subtitle_label.pack(pady=(0, 20))

        # 输出格式选择
        format_frame = ctk.CTkFrame(
            main_container,
            fg_color="#F8F8F8",
            corner_radius=8
        )
        format_frame.pack(fill="x", padx=15, pady=(10, 8))

        format_label = ctk.CTkLabel(
            format_frame,
            text="输出格式:",
            font=ctk.CTkFont(size=13, family="Microsoft YaHei"),
            text_color="#333333"
        )
        format_label.pack(side="left", padx=12, pady=8)

        self.format_var = ctk.StringVar(value="txt")

        txt_radio = ctk.CTkRadioButton(
            format_frame,
            text="TXT",
            variable=self.format_var,
            value="txt",
            command=self._on_format_change,
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            text_color="#333333"
        )
        txt_radio.pack(side="left", padx=8)

        word_radio = ctk.CTkRadioButton(
            format_frame,
            text="Word",
            variable=self.format_var,
            value="word",
            command=self._on_format_change,
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            text_color="#333333"
        )
        word_radio.pack(side="left", padx=8)

        # 日志区域
        log_frame = ctk.CTkFrame(
            main_container,
            fg_color="#FFFFFF",
            corner_radius=8,
            border_color="#1F6AA5",
            border_width=2
        )
        log_frame.pack(fill="both", expand=True, padx=15, pady=(10, 10))

        log_title = ctk.CTkLabel(
            log_frame,
            text="运行日志",
            font=ctk.CTkFont(size=13, weight="bold", family="Microsoft YaHei"),
            text_color="#333333",
            anchor="w"
        )
        log_title.pack(fill="x", padx=10, pady=(8, 5))

        self.log_text = ctk.CTkTextbox(
            log_frame,
            fg_color="#FFFFFF",
            text_color="#333333",
            font=ctk.CTkFont(size=11, family="Consolas"),
            corner_radius=6
        )
        self.log_text.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        # 状态栏
        self.status_label = ctk.CTkLabel(
            main_container,
            text="就绪",
            font=ctk.CTkFont(size=11, family="Microsoft YaHei"),
            text_color="#666666",
            anchor="w",
            fg_color="#F8F8F8",
            corner_radius=6
        )
        self.status_label.pack(fill="x", padx=15, pady=(10, 15))

        # 初始化日志
        self.log("=" * 60)
        self.log(" 资讯采集系统 V3.6 (独立版) 已启动")
        self.log("=" * 60)

        # 显示输出目录信息
        dirs = get_output_dirs()
        self.log(f" 输出目录: {dirs['base']}")
        self.log(f"   - 文档: {dirs['docs']}")
        self.log(f"   - 日志: {dirs['logs']}")
        self.log(f"   - 脚本: {dirs['scripts']}")
        self.log(f"   - 图片: {dirs['images']}")
        self.log(f"   - 其他: {dirs['other']}")

        # 显示统计
        self._show_stats()

    def _on_format_change(self):
        """格式改变事件"""
        self.output_format = self.format_var.get()
        self.log(f"输出格式已切换为: {self.output_format.upper()}")

    def _check_and_create_config_file(self):
        """检查并自动创建配置文件（小白友好化）"""
        import tkinter.messagebox as mb
        from tkinter import Tk

        # 获取配置文件路径
        if getattr(sys, 'frozen', False):
            # PyInstaller打包后的exe
            config_dir = os.path.dirname(sys.executable)
        else:
            # 开发环境
            config_dir = os.path.dirname(os.path.abspath(__file__))

        config_path = os.path.join(config_dir, 'prompts_config.yaml')

        # 检查配置文件是否存在
        if not os.path.exists(config_path):
            # 创建临时Tk窗口（用于弹窗）
            root = Tk()
            root.withdraw()  # 隐藏主窗口

            # 生成默认配置
            default_config = {
                'api': {
                    'api_url': 'https://open.bigmodel.cn/api/paas/v4/',
                    'api_key': '',  # 空密钥
                    'model': 'glm-4-flash',
                    'translation_temperature': 0.3,
                    'analysis_temperature': 0.7,
                    'translation_timeout': 30,
                    'analysis_timeout': 60
                },
                'translation': {
                    'en_to_zh_prompt': '请将以下英文内容翻译成中文，保持专业术语的准确性：\n\n{text}',
                    'enable_auto_skip': True
                },
                'analysis': {
                    'prompt_template': '请分析以下资讯内容...'
                }
            }

            # 写入配置文件
            try:
                import yaml
                with open(config_path, 'w', encoding='utf-8') as f:
                    yaml.safe_dump(default_config, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

                # 友好提示
                mb.showinfo(
                    "配置文件已自动生成",
                    "欢迎使用资讯采集系统 V36！\n\n"
                    "系统已为您自动生成了配置文件：\n"
                    f"{config_path}\n\n"
                    "【使用说明】\n"
                    "1. 基础功能（采集、分类、生成报告）无需配置API密钥\n"
                    "2. 如需使用AI智能分析功能，请填写API密钥后重启程序\n"
                    "3. 配置文件格式为YAML，可用记事本编辑\n\n"
                    "点击【确定】开始使用系统！",
                    parent=root
                )
            except Exception as e:
                mb.showwarning(
                    "配置文件生成失败",
                    f"自动生成配置文件时出错：{e}\n\n"
                    "程序仍可正常使用基础功能。",
                    parent=root
                )
            finally:
                root.destroy()

    def log(self, message: str):
        """日志输出"""
        timestamp = datetime.now().strftime('[%H:%M:%S]')
        log_entry = f"{timestamp} {message}"

        # 显示到界面（如果log_text已创建）
        if hasattr(self, 'log_text') and self.log_text:
            try:
                self.log_text.insert("end", f"{log_entry}\n")
                self.log_text.see("end")
                self.root.update()
            except:
                pass  # GUI还没准备好，静默失败

        # 保存到缓冲区
        if hasattr(self, 'log_buffer'):
            self.log_buffer.append(log_entry)

        # 如果GUI还没准备好，输出到控制台
        if not hasattr(self, 'log_text') or not self.log_text:
            print(log_entry)

    def save_log_to_file(self):
        """保存日志到文件"""
        try:
            dirs = get_output_dirs()
            log_dir = dirs['logs']

            # 日志文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            log_file = os.path.join(log_dir, f"运行日志_{timestamp}.txt")

            # 保存日志
            with open(log_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.log_buffer))

            self.log(f"[日志] 日志已保存: {log_file}")
        except Exception as e:
            self.log(f"[日志] 保存失败: {e}")

    def open_config(self):
        """打开配置文件"""
        config_path = os.path.join(os.path.dirname(__file__), 'classification_rules.yaml')
        if os.path.exists(config_path):
            os.startfile(config_path)
        else:
            self.log(f"配置文件不存在: {config_path}")

    def open_xiaohongshu_publisher(self):
        """打开小红书发布器（占位）"""
        self.log("小红书发布器功能待实现")

    def _run_xiaohongshu_task(self):
        """运行小红书任务（占位）"""
        pass

    def _show_xiaohongshu_preview(self):
        """显示小红书预览（占位）"""
        pass

    def _copy_to_clipboard(self):
        """复制到剪贴板（占位）"""
        pass

    def _save_xiaohongshu_content(self):
        """保存小红书内容（占位）"""
        pass

    def _open_xiaohongshu_browser(self):
        """打开小红书浏览器（占位）"""
        pass

    def _auto_publish_to_xiaohongshu(self):
        """自动发布到小红书（占位）"""
        pass

    def _execute_auto_publish(self):
        """执行自动发布（占位）"""
        pass

    def _wait_for_enter(self):
        """等待回车（占位）"""
        pass

    def _show_stats(self):
        """显示统计信息"""
        stats = self.db.get_stats()
        self.status_label.configure(
            text=f"总资讯: {stats['total']} | 待处理: {stats['pending']} | 已处理: {stats['processed']} | 已分析: {stats['analyzed']}"
        )

    def run_full(self):
        """全流程运行"""
        threading.Thread(target=self._run_full_task, daemon=True).start()

    def run_collect(self):
        """仅采集"""
        threading.Thread(target=self._run_collect_task, daemon=True).start()

    def run_process(self):
        """仅处理"""
        threading.Thread(target=self._run_process_task, daemon=True).start()

    def run_report(self):
        """生成报告"""
        threading.Thread(target=self._run_report_task, daemon=True).start()

    def _run_full_task(self):
        """全流程任务"""
        try:
            self.log("")
            self.log("=" * 60)
            self.log("[LAUNCH] 开始全流程任务")
            self.log("=" * 60)
            self.log("")

            # 步骤1：采集
            self.log("[步骤 1/2] [RSS] 采集资讯...")
            stats = self.collector.fetch_all()

            # 步骤2：处理本次采集的新资讯
            self.log("")
            self.log("[步骤 2/2] [SETTINGS] 处理资讯...")
            self.processor.process_all(item_ids=stats['item_ids'])

            # 完成
            self.log("")
            self.log("=" * 60)
            self.log("[OK] 全流程任务完成！")
            self.log("=" * 60)

            # 保存日志
            self.save_log_to_file()

            # 生成报告
            self._run_report_task()

            # 更新统计
            self._show_stats()

        except Exception as e:
            self.log(f"[ERROR] 任务失败: {e}")
            self.log(traceback.format_exc())

    def _run_collect_task(self):
        """采集任务"""
        try:
            self.log("")
            self.log("[采集] 开始采集资讯...")
            stats = self.collector.fetch_all()
            self.log(f"[采集] 完成，共保存 {stats['saved']} 条资讯")
            self._show_stats()

        except Exception as e:
            self.log(f"[ERROR] 采集失败: {e}")

    def _run_process_task(self):
        """处理任务"""
        try:
            self.log("")
            self.log("[处理] 开始处理资讯...")
            count = self.processor.process_all()
            self.log(f"[处理] 完成，共处理 {count} 条资讯")
            self._show_stats()

        except Exception as e:
            self.log(f"[ERROR] 处理失败: {e}")

    def _run_report_task(self):
        """报告任务"""
        try:
            # 获取今天的日期字符串（用于过滤和文件名）
            today = datetime.now().strftime('%Y-%m-%d')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')  # 精确到分钟的时间戳

            # 获取截止到今天的所有待报告资讯（按处理时间过滤，包含历史遗留的未报告资讯）
            items = self.db.get_processed_items(date=today, use_processed_time=True)

            # 空数据保护：如果没有待生成的资讯（status='processed'），直接返回
            if not items:
                self.log(f"[报告] 目前没有待生成的资讯（status='processed'），跳过任务")
                return

            # 生成报告（可能包含历史遗留的待报告资讯）
            self.log(f"[报告] 发现共 {len(items)} 条待报告资讯（含历史遗留），正在生成报告...")

            # 获取输出目录
            dirs = get_output_dirs()
            doc_dir = dirs['docs']

            # 生成TXT报告（文件名包含时间戳，避免覆盖）
            report_filename = f"资讯深度分析报告_{today}_{timestamp}.txt"
            report_path = os.path.join(doc_dir, report_filename)
            report_content = self.processor._generate_report(items)

            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)

            self.log(f"[报告] TXT报告已保存: {report_path}")

            # ========== [新增] 同步生成CSV报告 ==========
            try:
                if CSV_AVAILABLE:
                    # 生成CSV文件路径（与TXT同名）
                    csv_filename = report_filename.replace('.txt', '.csv')
                    csv_path = os.path.join(doc_dir, csv_filename)

                    # 调用CSV导出模块
                    csv_exporter.generate_csv_sync(report_path, items)

                    self.log(f"[报告] CSV报告已同步保存: {csv_path}")
                else:
                    self.log(f"[报告] CSV导出模块未加载，跳过CSV生成")

            except Exception as csv_error:
                # CSV生成失败不影响TXT报告
                self.log(f"[报告]  CSV生成失败（TXT报告已正常保存）: {csv_error}")
                self.log(traceback.format_exc())

            # 生成Word报告（如果需要）
            if self.output_format == "word":
                word_filename = f"资讯深度分析报告_{today}_{timestamp}.docx"
                word_path = os.path.join(doc_dir, word_filename)
                self.processor.export_to_word(word_path, items)
                self.log(f"[报告] Word报告已保存: {word_path}")

            self.log(f"[报告] 报告生成完成！文件名: {report_filename}")

            # 标记这些资讯为已报告（避免重复生成报告）
            item_ids = [item['id'] for item in items]
            self.db.mark_as_processed(item_ids)
            self.log(f"[报告] [OK] 报告生成完毕，{len(item_ids)} 条资讯状态已更新为 'reported'")

        except Exception as e:
            self.log(f"[ERROR] 报告生成失败: {e}")
            self.log(traceback.format_exc())

    def extract_source_code(self):
        """提取源代码到当前目录"""
        try:
            import shutil

            # 如果是PyInstaller打包的exe
            if getattr(sys, 'frozen', False):
                # 当前exe目录
                exe_dir = os.path.dirname(os.path.abspath(sys.executable))

                # PyInstaller临时目录
                if '_MEIPASS' in os.environ:
                    temp_dir = os.environ['_MEIPASS']

                    # 源文件路径
                    temp_py = os.path.join(temp_dir, 'source', 'main_gui_standalone.py')
                    temp_spec = os.path.join(temp_dir, 'source', 'NewsCollector_V34.spec')

                    # 目标路径
                    target_py = os.path.join(exe_dir, 'main_gui_standalone.py')
                    target_spec = os.path.join(exe_dir, 'NewsCollector_V34.spec')

                    # 提取源文件
                    if os.path.exists(temp_py):
                        shutil.copy2(temp_py, target_py)
                        self.log(f"[源代码] [OK] 已提取源文件: {target_py}")
                    else:
                        self.log("[源代码] [X] 源文件不存在")

                    # 提取配置文件
                    if os.path.exists(temp_spec):
                        shutil.copy2(temp_spec, target_spec)
                        self.log(f"[源代码] [OK] 已提取配置文件: {target_spec}")
                    else:
                        self.log("[源代码] [X] 配置文件不存在")

                    # 检查classification_rules.yaml
                    yaml_file = os.path.join(exe_dir, 'classification_rules.yaml')
                    if not os.path.exists(yaml_file):
                        temp_yaml = os.path.join(temp_dir, 'classification_rules.yaml')
                        if os.path.exists(temp_yaml):
                            shutil.copy2(temp_yaml, yaml_file)
                            self.log(f"[源代码] [OK] 已提取配置文件: {yaml_file}")

                    self.log("[源代码] [OK] 提取完成！现在可以修改源代码了")
                    self.log("[源代码] 提示: 修改后运行 pyinstaller NewsCollector_V34.spec 重新打包")

                else:
                    self.log("[源代码] [X] 无法访问临时目录")
            else:
                self.log("[源代码]  当前是开发模式，源代码已存在")

        except Exception as e:
            self.log(f"[源代码] [X] 提取失败: {e}")
            self.log(traceback.format_exc())

    def open_settings(self):
        """打开设置窗口"""
        # 创建设置窗口
        settings_window = ctk.CTkToplevel(self.root)
        settings_window.title("API设置")
        settings_window.geometry("500x480")
        settings_window.resizable(False, False)
        settings_window.transient(self.root)
        settings_window.grab_set()

        # 设置窗口背景色
        settings_window.configure(fg_color="#F2F2F2")

        # 主容器
        container = ctk.CTkFrame(
            settings_window,
            fg_color="#FFFFFF",
            corner_radius=10
        )
        container.pack(fill="both", expand=True, padx=20, pady=20)

        # 标题
        title_label = ctk.CTkLabel(
            container,
            text="API配置设置",
            font=ctk.CTkFont(size=22, weight="bold", family="Microsoft YaHei"),
            text_color="#333333"
        )
        title_label.pack(pady=(15, 20))

        # 当前API配置显示
        self._load_current_api_config()

        # API地址
        api_url_label = ctk.CTkLabel(
            container,
            text="API地址:",
            font=ctk.CTkFont(size=13, family="Microsoft YaHei"),
            text_color="#333333",
            anchor="w"
        )
        api_url_label.pack(fill="x", pady=(10, 5), padx=15)
        api_url_entry = ctk.CTkEntry(
            container,
            placeholder_text="",
            height=40,
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            fg_color="#F8F8F8",
            border_color="#1F6AA5",
            border_width=2,
            corner_radius=6
        )
        api_url_entry.pack(fill="x", pady=5, padx=15)
        api_url_entry.insert(0, self.current_api_url)

        # API密钥
        api_key_label = ctk.CTkLabel(
            container,
            text="API密钥:",
            font=ctk.CTkFont(size=13, family="Microsoft YaHei"),
            text_color="#333333",
            anchor="w"
        )
        api_key_label.pack(fill="x", pady=(10, 5), padx=15)
        api_key_entry = ctk.CTkEntry(
            container,
            placeholder_text="",
            show="*",
            height=40,
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            fg_color="#F8F8F8",
            border_color="#1F6AA5",
            border_width=2,
            corner_radius=6
        )
        api_key_entry.pack(fill="x", pady=5, padx=15)
        api_key_entry.insert(0, self.current_api_key)

        # 模型名称
        model_label = ctk.CTkLabel(
            container,
            text="模型名称:",
            font=ctk.CTkFont(size=13, family="Microsoft YaHei"),
            text_color="#333333",
            anchor="w"
        )
        model_label.pack(fill="x", pady=(10, 5), padx=15)
        model_entry = ctk.CTkEntry(
            container,
            placeholder_text="",
            height=40,
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            fg_color="#F8F8F8",
            border_color="#1F6AA5",
            border_width=2,
            corner_radius=6
        )
        model_entry.pack(fill="x", pady=5, padx=15)
        model_entry.insert(0, self.current_model)

        # 提示信息
        info_label = ctk.CTkLabel(
            container,
            text="修改后需要重启程序才能生效",
            font=ctk.CTkFont(size=12, family="Microsoft YaHei"),
            justify="center",
            text_color="#FF6B35"
        )
        info_label.pack(fill="x", pady=20, padx=15)

        # 按钮区域
        button_frame = ctk.CTkFrame(container, fg_color="transparent")
        button_frame.pack(fill="x", padx=15, pady=15)

        def save_settings():
            """保存设置"""
            api_url = api_url_entry.get().strip()
            api_key = api_key_entry.get().strip()
            model = model_entry.get().strip()

            if not api_url or not api_key or not model:
                self.log("[ERROR] 配置不完整，请填写所有字段")
                return

            try:
                self._save_api_config(api_url, api_key, model)
                self.log("[OK] API配置已保存，请重启程序使配置生效")
                settings_window.destroy()
                messagebox.showinfo(
                    "配置已保存",
                    "API配置已成功保存！\n\n请关闭程序后重新启动，新配置才会生效。"
                )
            except Exception as e:
                self.log(f"[ERROR] 保存配置失败: {e}")
                messagebox.showerror("保存失败", f"保存配置时出错：{e}")

        def test_connection():
            """测试API连接"""
            api_url = api_url_entry.get().strip()
            api_key = api_key_entry.get().strip()
            model = model_entry.get().strip()

            if not api_url or not api_key or not model:
                messagebox.showwarning("配置不完整", "请先填写所有字段")
                return

            try:
                import requests
                test_url = api_url.rstrip('/') + '/chat/completions'
                response = requests.post(
                    test_url,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": "Hi"}],
                        "max_tokens": 1
                    },
                    timeout=10
                )

                if response.status_code == 200:
                    messagebox.showinfo("测试成功", "[OK] API连接测试成功！\n\n配置正确，可以正常使用。")
                elif response.status_code == 401:
                    messagebox.showerror("测试失败", "[ERROR] API密钥无效\n\n请检查API密钥是否正确")
                else:
                    error_msg = response.text[:200] if response.text else "未知错误"
                    messagebox.showerror("测试失败", f"[ERROR] API测试失败\n\n状态码：{response.status_code}\n错误：{error_msg}")
            except Exception as e:
                messagebox.showerror("测试失败", f"[ERROR] 连接测试失败\n\n错误：{str(e)}")

        # 按钮样式配置
        button_style = {
            "fg_color": "#1F6AA5",
            "hover_color": "#3B8ED0",
            "text_color": "#FFFFFF",
            "corner_radius": 8,
            "font": ctk.CTkFont(size=13, weight="bold", family="Microsoft YaHei"),
            "height": 38
        }

        save_btn = ctk.CTkButton(
            button_frame,
            text="[SAVE] 保存配置",
            command=save_settings,
            width=150,
            **button_style
        )
        save_btn.pack(side="left", padx=5)

        test_btn = ctk.CTkButton(
            button_frame,
            text="[TEST] 测试连接",
            command=test_connection,
            width=150,
            **button_style
        )
        test_btn.pack(side="left", padx=5)

        cancel_btn = ctk.CTkButton(
            button_frame,
            text="[ERROR] 取消",
            command=settings_window.destroy,
            width=150,
            fg_color="#999999",
            hover_color="#BBBBBB",
            **{k: v for k, v in button_style.items() if k not in ["fg_color", "hover_color"]}
        )
        cancel_btn.pack(side="left", padx=5)

    def _load_current_api_config(self):
        """加载当前API配置到实例变量"""
        try:
            # 优先从外部配置文件读取
            if getattr(sys, 'frozen', False):
                # PyInstaller打包后的exe：使用exe所在目录
                exe_dir = os.path.dirname(sys.executable)
                config_path = os.path.join(exe_dir, 'prompts_config.yaml')
            else:
                # 开发环境：使用脚本所在目录
                config_path = os.path.join(os.path.dirname(__file__), 'prompts_config.yaml')

            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
                    api_config = config.get('api', {})
                    self.current_api_url = api_config.get('api_url', 'https://open.bigmodel.cn/api/paas/v4/')
                    self.current_api_key = api_config.get('api_key', '')
                    self.current_model = api_config.get('model', 'glm-4-flash')
            else:
                # 使用默认值
                self.current_api_url = 'https://open.bigmodel.cn/api/paas/v4/'
                self.current_api_key = ''
                self.current_model = 'glm-4-flash'
        except Exception as e:
            self.log(f"[WARNING] 加载API配置失败: {e}")
            self.current_api_url = 'https://open.bigmodel.cn/api/paas/v4/'
            self.current_api_key = ''
            self.current_model = 'glm-4-flash'

    def _save_api_config(self, api_url: str, api_key: str, model: str):
        """保存API配置到prompts_config.yaml"""
        try:
            # 确定配置文件路径
            if getattr(sys, 'frozen', False):
                exe_dir = os.path.dirname(sys.executable)
                config_path = os.path.join(exe_dir, 'prompts_config.yaml')
            else:
                config_path = os.path.join(os.path.dirname(__file__), 'prompts_config.yaml')

            # 读取现有配置
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = yaml.safe_load(f)
            else:
                # 配置文件不存在，创建带有默认结构的模板
                config = {
                    'api': {
                        'api_url': 'https://open.bigmodel.cn/api/paas/v4/',
                        'api_key': '',
                        'model': 'glm-4-flash',
                        'translation_temperature': 0.3,
                        'analysis_temperature': 0.7,
                        'translation_timeout': 30,
                        'analysis_timeout': 60
                    },
                    'translation': {
                        'en_to_zh_prompt': '请将以下英文内容翻译成中文，保持专业术语的准确性：\n\n{text}',
                        'enable_auto_skip': True
                    },
                    'analysis': {
                        'prompt_template': '请分析以下资讯：\n\n标题: {title}\n\n内容: {content}',
                        'max_content_length': 1000,
                        'skip_keywords': ['[跳过', '上市促销']
                    }
                }
                self.log(f"[OK] 创建新的配置文件模板: {config_path}")

            # 更新API配置
            if 'api' not in config:
                config['api'] = {}
            config['api']['api_url'] = api_url
            config['api']['api_key'] = api_key
            config['api']['model'] = model

            # 保存配置
            with open(config_path, 'w', encoding='utf-8') as f:
                yaml.dump(config, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

            self.log(f"[OK] 配置已保存到: {config_path}")
        except Exception as e:
            raise Exception(f"保存配置失败: {e}")

    def open_rss_manager(self):
        """打开RSS源管理窗口"""
        # 创建RSS源管理窗口
        rss_window = ctk.CTkToplevel(self.root)
        rss_window.title("RSS源管理")
        rss_window.geometry("800x600")
        rss_window.transient(self.root)
        rss_window.grab_set()

        # 设置窗口背景色
        rss_window.configure(fg_color="#F2F2F2")

        # 标题
        title_label = ctk.CTkLabel(
            rss_window,
            text="RSS源管理",
            font=ctk.CTkFont(size=22, weight="bold", family="Microsoft YaHei"),
            text_color="#333333"
        )
        title_label.pack(pady=(15, 10))

        # 主容器
        main_container = ctk.CTkFrame(
            rss_window,
            fg_color="#FFFFFF",
            corner_radius=10
        )
        main_container.pack(fill="both", expand=True, padx=20, pady=(10, 20))

        # 左侧：RSS源列表
        left_frame = ctk.CTkFrame(main_container)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

        ctk.CTkLabel(left_frame, text="RSS源列表：", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 5))

        # 创建列表框架（带滚动条）
        list_frame = ctk.CTkScrollableFrame(left_frame, width=500, height=400)
        list_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # 右侧：添加新源
        right_frame = ctk.CTkFrame(main_container)
        right_frame.pack(side="right", fill="y", padx=(10, 0))

        ctk.CTkLabel(right_frame, text="添加新RSS源", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 15))

        # RSS URL
        ctk.CTkLabel(right_frame, text="RSS地址:").pack(anchor="w", padx=10)
        url_entry = ctk.CTkEntry(right_frame, placeholder_text="", width=250, height=35)
        url_entry.pack(padx=10, pady=5)

        # 名称
        ctk.CTkLabel(right_frame, text="名称:").pack(anchor="w", padx=10)
        name_entry = ctk.CTkEntry(right_frame, placeholder_text="", width=250, height=35)
        name_entry.pack(padx=10, pady=5)

        # 分类
        ctk.CTkLabel(right_frame, text="分类:").pack(anchor="w", padx=10)
        category_var = ctk.StringVar(value="ai_tech")
        category_menu = ctk.CTkOptionMenu(
            right_frame,
            variable=category_var,
            values=["ai_tech", "medical", "news", "movie", "gaming"],
            width=250,
            height=35
        )
        category_menu.pack(padx=10, pady=5)

        # 添加按钮
        def add_rss_source():
            url = url_entry.get().strip()
            name = name_entry.get().strip()
            category = category_var.get()

            if not url:
                messagebox.showwarning("输入错误", "请输入RSS地址")
                return

            # [OK] 资源回收：确保cursor在所有分支都能关闭
            cursor = self.db.conn.cursor()
            try:
                # 检查URL是否已存在
                cursor.execute('SELECT id FROM rss_sources WHERE url = ?', (url,))
                if cursor.fetchone():
                    messagebox.showerror("添加失败", "该RSS源已存在！")
                    return

                # 保存到数据库
                from datetime import datetime
                now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                cursor.execute('''
                    INSERT INTO rss_sources (url, name, category, source_type, is_enabled, created_at)
                    VALUES (?, ?, ?, ?, 1, ?)
                ''', (url, name or url, category, category, now))
                self.db.conn.commit()

                self.log(f"[OK] 已添加RSS源: {name or url}")
                url_entry.delete(0, 'end')
                name_entry.delete(0, 'end')

                # 刷新列表
                refresh_rss_list()
                messagebox.showinfo("添加成功", "RSS源已成功添加！")

            except Exception as e:
                messagebox.showerror("添加失败", f"添加RSS源失败：{str(e)}")
            finally:
                cursor.close()

        # 测试按钮
        def test_rss_source():
            url = url_entry.get().strip()
            if not url:
                messagebox.showwarning("输入错误", "请输入RSS地址进行测试")
                return

            try:
                import requests
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    'Accept': 'application/rss+xml, application/xml, text/xml, */*',
                }
                response = requests.head(url, headers=headers, timeout=10, verify=False)

                if response.status_code in [200, 301, 302]:
                    messagebox.showinfo("测试成功", "[OK] RSS源可用！\n\n可以添加此源。")
                else:
                    messagebox.showerror("测试失败", f"[ERROR] RSS源不可用\n\n状态码：{response.status_code}")
            except Exception as e:
                messagebox.showerror("测试失败", f"[ERROR] 连接失败\n\n错误：{str(e)}")

        # 按钮区域
        ctk.CTkButton(right_frame, text="[ADD] 添加RSS源", command=add_rss_source, width=250, height=40, fg_color="green").pack(pady=10)
        ctk.CTkButton(right_frame, text="[TEST] 测试连接", command=test_rss_source, width=250, height=40).pack(pady=10)

        # 刷新列表的函数
        def refresh_rss_list():
            # 清空列表
            for widget in list_frame.winfo_children():
                widget.destroy()

            # [OK] 资源回收：立即获取数据并关闭cursor
            cursor = self.db.conn.cursor()
            try:
                cursor.execute('''
                    SELECT id, url, name, category, is_enabled
                    FROM rss_sources
                    ORDER BY category, id
                ''')
                sources = cursor.fetchall()
            finally:
                cursor.close()

            if not sources:
                no_data_label = ctk.CTkLabel(list_frame, text="暂无RSS源", text_color="gray")
                no_data_label.pack(pady=20)
                return

            # 显示RSS源
            for idx, (source_id, url, name, category, is_enabled) in enumerate(sources, start=1):
                item_frame = ctk.CTkFrame(list_frame)
                item_frame.pack(fill="x", pady=5, padx=5)

                # 信息行（添加编号）
                info_text = f"{idx}. [{category}] {name or url}"
                info_label = ctk.CTkLabel(item_frame, text=info_text, anchor="w", font=ctk.CTkFont(size=12))
                info_label.pack(fill="x", padx=10, pady=(5, 0))

                # URL行
                url_label = ctk.CTkLabel(item_frame, text=url, anchor="w", text_color="gray", font=ctk.CTkFont(size=10))
                url_label.pack(fill="x", padx=10, pady=(0, 5))

                # 按钮行
                btn_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
                btn_frame.pack(fill="x", padx=10, pady=(0, 5))

                # 删除按钮
                # [OK] 修复：在嵌套函数中创建新的cursor，避免闭包问题
                def delete_source(sid=source_id, name=name or url):
                    if messagebox.askyesno("确认删除", f"确定要删除RSS源吗？\n\n{name}"):
                        delete_cursor = self.db.conn.cursor()
                        try:
                            delete_cursor.execute('DELETE FROM rss_sources WHERE id = ?', (sid,))
                            self.db.conn.commit()
                            self.log(f"[DELETE] 已删除RSS源: {name}")
                            refresh_rss_list()
                            messagebox.showinfo("删除成功", "RSS源已删除")
                        except Exception as e:
                            messagebox.showerror("删除失败", f"删除失败：{str(e)}")
                        finally:
                            delete_cursor.close()

                ctk.CTkButton(btn_frame, text="[DELETE] 删除", command=delete_source, width=80, height=30, fg_color="red").pack(side="right")

        # 初始加载列表
        refresh_rss_list()

    def open_keyword_manager(self):
        """打开自定义关键字管理窗口"""
        # 创建关键字管理窗口
        keyword_window = ctk.CTkToplevel(self.root)
        keyword_window.title("自定义关键字管理")
        keyword_window.geometry("800x600")
        keyword_window.transient(self.root)
        keyword_window.grab_set()

        # 设置窗口背景色
        keyword_window.configure(fg_color="#F2F2F2")

        # 标题
        title_label = ctk.CTkLabel(
            keyword_window,
            text="自定义关键字管理",
            font=ctk.CTkFont(size=22, weight="bold", family="Microsoft YaHei"),
            text_color="#333333"
        )
        title_label.pack(pady=(15, 5))

        # 说明标签
        info_label = ctk.CTkLabel(
            keyword_window,
            text="设置核心关键字后，程序只处理包含这些关键字的资讯",
            text_color="#666666",
            font=ctk.CTkFont(size=11, family="Microsoft YaHei")
        )
        info_label.pack(pady=(0, 10))

        # 主容器
        main_container = ctk.CTkFrame(
            keyword_window,
            fg_color="#FFFFFF",
            corner_radius=10
        )
        main_container.pack(fill="both", expand=True, padx=20, pady=(10, 20))

        # 左侧：关键字列表
        left_frame = ctk.CTkFrame(main_container)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

        ctk.CTkLabel(left_frame, text="已添加的关键字：", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 5))

        # 创建列表框架（带滚动条）
        list_frame = ctk.CTkScrollableFrame(left_frame, width=500, height=400)
        list_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # 右侧：添加新关键字
        right_frame = ctk.CTkFrame(main_container)
        right_frame.pack(side="right", fill="y", padx=(10, 0))

        ctk.CTkLabel(right_frame, text="添加新关键字", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=(10, 15))

        # 关键字输入
        ctk.CTkLabel(right_frame, text="关键字:").pack(anchor="w", padx=10)
        keyword_entry = ctk.CTkEntry(right_frame, placeholder_text="例如：AI、人工智能、医疗...", width=250, height=35)
        keyword_entry.pack(padx=10, pady=5)

        # 描述
        ctk.CTkLabel(right_frame, text="描述说明:").pack(anchor="w", padx=10)
        description_entry = ctk.CTkEntry(right_frame, placeholder_text="可选：说明此关键字的用途", width=250, height=35)
        description_entry.pack(padx=10, pady=5)

        # 分类
        ctk.CTkLabel(right_frame, text="分类:").pack(anchor="w", padx=10)
        category_var = ctk.StringVar(value="default")
        category_menu = ctk.CTkOptionMenu(
            right_frame,
            variable=category_var,
            values=["default", "tech", "medical", "business", "other"],
            width=250,
            height=35
        )
        category_menu.pack(padx=10, pady=5)

        # 添加按钮
        def add_keyword():
            keyword = keyword_entry.get().strip()
            description = description_entry.get().strip()
            category = category_var.get()

            if not keyword:
                messagebox.showwarning("输入错误", "请输入关键字")
                return

            # 添加到数据库
            result = self.db.add_custom_keyword(keyword, description, category)

            if result == -1:
                messagebox.showerror("添加失败", "该关键字已存在！")
            elif result > 0:
                self.log(f"[OK] 已添加关键字: {keyword}")
                keyword_entry.delete(0, 'end')
                description_entry.delete(0, 'end')

                # 刷新列表
                refresh_keyword_list()
                messagebox.showinfo("添加成功", "关键字已成功添加！")
            else:
                messagebox.showerror("添加失败", "添加失败，请重试")

        # 按钮区域
        ctk.CTkButton(right_frame, text="[ADD] 添加关键字", command=add_keyword, width=250, height=40, fg_color="green").pack(pady=10)

        # 刷新列表的函数
        def refresh_keyword_list():
            # 清空列表
            for widget in list_frame.winfo_children():
                widget.destroy()

            # 获取关键字列表
            keywords = self.db.get_custom_keywords(enabled_only=False)

            if not keywords:
                no_data_label = ctk.CTkLabel(list_frame, text="暂无关键字", text_color="gray")
                no_data_label.pack(pady=20)
                return

            # 显示关键字
            for idx, kw in enumerate(keywords, start=1):
                item_frame = ctk.CTkFrame(list_frame)
                item_frame.pack(fill="x", pady=5, padx=5)

                # 状态指示
                status_color = "green" if kw['is_enabled'] else "gray"
                status_text = "[OK]" if kw['is_enabled'] else "[X]"

                # 信息行
                info_text = f"{idx}. {kw['keyword']}"
                info_label = ctk.CTkLabel(item_frame, text=info_text, anchor="w", font=ctk.CTkFont(size=12))
                info_label.pack(fill="x", padx=10, pady=(5, 0))

                # 描述行
                if kw['description']:
                    desc_label = ctk.CTkLabel(
                        item_frame,
                        text=f"描述: {kw['description']}",
                        anchor="w",
                        text_color="gray",
                        font=ctk.CTkFont(size=10)
                    )
                    desc_label.pack(fill="x", padx=10, pady=(0, 5))

                # 按钮行
                btn_frame = ctk.CTkFrame(item_frame, fg_color="transparent")
                btn_frame.pack(fill="x", padx=10, pady=(0, 5))

                # 启用/禁用按钮
                def toggle_keyword(kwid=kw['id'], current_status=kw['is_enabled'], kwname=kw['keyword']):
                    new_status = not current_status
                    self.db.update_custom_keyword(kwid, is_enabled=1 if new_status else 0)
                    self.log(f"{'[OK] 已启用' if new_status else '[DISABLED] 已禁用'} 关键字: {kwname}")
                    refresh_keyword_list()

                toggle_text = "[DISABLED] 禁用" if kw['is_enabled'] else "[OK] 启用"
                toggle_color = "orange" if kw['is_enabled'] else "green"
                ctk.CTkButton(
                    btn_frame,
                    text=toggle_text,
                    command=toggle_keyword,
                    width=80,
                    height=30,
                    fg_color=toggle_color
                ).pack(side="left", padx=(0, 5))

                # 删除按钮
                def delete_keyword(kwid=kw['id'], kwname=kw['keyword']):
                    if messagebox.askyesno("确认删除", f"确定要删除关键字吗？\n\n{kwname}"):
                        if self.db.delete_custom_keyword(kwid):
                            self.log(f"[DELETE] 已删除关键字: {kwname}")
                            refresh_keyword_list()
                            messagebox.showinfo("删除成功", "关键字已删除")
                        else:
                            messagebox.showerror("删除失败", "删除失败，请重试")

                ctk.CTkButton(btn_frame, text="[DELETE] 删除", command=delete_keyword, width=80, height=30, fg_color="red").pack(side="right")

        # 初始加载列表
        refresh_keyword_list()

    def cleanup_irrelevant(self):
        """清理不匹配自定义关键字的待处理资讯"""
        try:
            self.log("")
            self.log("=" * 60)
            self.log("[CLEANUP] 开始清理不相关资讯...")
            self.log("=" * 60)

            # 检查是否有启用的自定义关键字
            custom_keywords = self.db.get_custom_keywords(enabled_only=True)
            if not custom_keywords:
                self.log("[CLEANUP] 未启用自定义关键字，无需清理")
                self.log("=" * 60)
                self.log("")
                return

            # 显示当前启用的关键字
            keyword_list = [kw['keyword'] for kw in custom_keywords]
            self.log(f"[CLEANUP] 当前启用的关键字: {', '.join(keyword_list)}")

            # 执行清理
            count = self.db.cleanup_irrelevant_items()

            if count > 0:
                self.log(f"[CLEANUP] 已将 {count} 条不相关资讯标记为 'ignored' 状态")
            else:
                self.log("[CLEANUP] 没有需要清理的资讯")

            # 更新统计显示
            self._show_stats()

            self.log("=" * 60)
            self.log("[OK] 清理完成！")
            self.log("=" * 60)
            self.log("")

        except Exception as e:
            self.log(f"[ERROR] 清理失败: {e}")
            self.log(traceback.format_exc())

    def run(self):
        """运行应用"""
        self.root.mainloop()

# ============================================
# 主程序入口
# ============================================

if __name__ == "__main__":
    # 提取源代码（如果不存在）
    extract_source_if_needed()

    # 启动应用
    app = NewsAppGUI()
    app.run()
